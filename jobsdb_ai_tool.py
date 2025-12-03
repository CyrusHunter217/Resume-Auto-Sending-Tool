#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
JobsDB 自动求职工具
功能：自动搜索岗位、生成定制简历、自动投递
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import json
import os
import sys
import threading
import time
import random
import re
from datetime import datetime
from urllib.parse import urljoin, urlparse
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
try:
    from webdriver_manager.chrome import ChromeDriverManager
    WEBDRIVER_MANAGER_AVAILABLE = True
except ImportError:
    WEBDRIVER_MANAGER_AVAILABLE = False
from fpdf import FPDF
import pandas as pd
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
import webbrowser

# 获取资源路径（支持打包后的exe）
def get_resource_path(relative_path):
    """获取资源文件的绝对路径，支持打包后的exe"""
    if getattr(sys, 'frozen', False):
        # 打包后的exe
        base_path = sys._MEIPASS
    else:
        # 开发环境
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


def detect_chrome_user_data_dir():
    """
    自动检测Chrome用户数据目录路径
    返回: (路径, 是否检测成功, 提示信息)
    """
    try:
        if sys.platform == "win32":
            # Windows系统
            username = os.getenv('USERNAME') or os.getenv('USER')
            if username:
                default_path = os.path.join(
                    os.getenv('LOCALAPPDATA', ''),
                    'Google', 'Chrome', 'User Data'
                )
                # 检查路径是否存在
                if os.path.exists(default_path):
                    return default_path, True, "已自动检测到Chrome路径"
                else:
                    # 尝试另一个可能的路径
                    alt_path = os.path.join(
                        'C:', 'Users', username, 'AppData', 'Local', 'Google', 'Chrome', 'User Data'
                    )
                    if os.path.exists(alt_path):
                        return alt_path, True, "已自动检测到Chrome路径"
            return "", False, "未检测到Chrome，将使用新浏览器窗口（首次需手动登录）"
        
        elif sys.platform == "darwin":
            # macOS系统
            username = os.getenv('USER')
            if username:
                default_path = os.path.join(
                    os.path.expanduser('~'),
                    'Library', 'Application Support', 'Google', 'Chrome'
                )
                if os.path.exists(default_path):
                    return default_path, True, "已自动检测到Chrome路径"
            return "", False, "未检测到Chrome，将使用新浏览器窗口（首次需手动登录）"
        
        else:
            # Linux等其他系统
            username = os.getenv('USER')
            if username:
                default_path = os.path.join(
                    os.path.expanduser('~'),
                    '.config', 'google-chrome'
                )
                if os.path.exists(default_path):
                    return default_path, True, "已自动检测到Chrome路径"
            return "", False, "未检测到Chrome，将使用新浏览器窗口（首次需手动登录）"
    
    except Exception as e:
        return "", False, f"检测失败: {str(e)}，将使用新浏览器窗口"


def scan_chrome_profiles(user_data_dir):
    """
    扫描Chrome用户数据目录下的所有配置文件
    返回: 配置文件名称列表，例如 ['Default', 'Profile 1', 'Profile 2']
    """
    profiles = []
    if not user_data_dir or not os.path.exists(user_data_dir):
        return profiles
    
    try:
        # 扫描User Data目录下的所有Profile文件夹
        for item in os.listdir(user_data_dir):
            item_path = os.path.join(user_data_dir, item)
            if os.path.isdir(item_path):
                # Chrome配置文件通常以"Default"或"Profile X"命名
                if item == "Default" or item.startswith("Profile "):
                    # 检查是否是有效的配置文件（通常包含Preferences文件）
                    prefs_file = os.path.join(item_path, "Preferences")
                    if os.path.exists(prefs_file):
                        profiles.append(item)
        
        # 确保Default在第一位
        if "Default" in profiles:
            profiles.remove("Default")
            profiles.insert(0, "Default")
        
        return profiles
    except Exception as e:
        print(f"扫描配置文件失败: {e}")
        return profiles


def check_chrome_running():
    """
    检查Chrome是否正在运行
    返回: (是否运行, 进程数量)
    """
    try:
        if sys.platform == "win32":
            # Windows系统
            import subprocess
            result = subprocess.run(
                ['tasklist', '/FI', 'IMAGENAME eq chrome.exe'],
                capture_output=True,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
            )
            # 统计chrome.exe进程数量（排除标题行）
            lines = result.stdout.strip().split('\n')
            count = len([line for line in lines if 'chrome.exe' in line.lower()]) - 1
            return count > 0, count
        elif sys.platform == "darwin":
            # macOS系统
            import subprocess
            result = subprocess.run(
                ['pgrep', '-f', 'Google Chrome'],
                capture_output=True
            )
            count = len(result.stdout.decode().strip().split('\n')) if result.stdout else 0
            return count > 0, count
        else:
            # Linux系统
            import subprocess
            result = subprocess.run(
                ['pgrep', '-f', 'chrome'],
                capture_output=True
            )
            count = len(result.stdout.decode().strip().split('\n')) if result.stdout else 0
            return count > 0, count
    except Exception as e:
        print(f"检查Chrome进程失败: {e}")
        return False, 0


class ResumeGeneratorApp:
    """主应用程序类"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("JobsDB 自动求职工具")
        self.root.geometry("1000x700")
        
        # 初始化变量（需要在load_config之前设置默认值）
        self.language = "zh"  # 默认中文
        self.texts = self.get_texts("zh")
        
        # 初始化变量
        self.is_auto_running = False
        self.is_paused = False
        self.pause_event = threading.Event()
        self.pause_event.set()  # 初始状态为运行
        
        # 存储GUI元素引用（用于语言切换）
        self.ui_labels = {}
        self.ui_buttons = {}
        self.ui_frames = {}
        
        # 简历同步标志（防止循环同步）
        self._syncing_resume = False
        
        # 加载配置
        self.config_file = "config.json"
        self.load_config()
        
        # 恢复语言设置
        if 'language' in self.config:
            self.language = self.config['language']
            self.texts = self.get_texts(self.language)
        
        # 创建界面
        self.create_widgets()
        
        # 加载简历缓存（在界面创建后）
        self.load_resume_from_cache()
        
        # 设置自动保存
        self.setup_auto_save()
    
    def load_resume_from_cache(self):
        """从缓存加载简历"""
        # 如果配置中有简历内容，直接使用
        if 'resume_content' in self.config and self.config['resume_content']:
            if hasattr(self, 'resume_text_init'):
                current_content = self.resume_text_init.get("1.0", tk.END).strip()
                if not current_content:
                    self.resume_text_init.insert("1.0", self.config['resume_content'])
        else:
            # 尝试从缓存文件加载
            if hasattr(self, 'resume_text_init'):
                cached_resume = self.load_resume_cache()
                if cached_resume:
                    current_content = self.resume_text_init.get("1.0", tk.END).strip()
                    if not current_content:
                        self.resume_text_init.insert("1.0", cached_resume)
                        # 同时保存到配置
                        self.config['resume_content'] = cached_resume
                        self.save_config()
    
    def get_texts(self, lang="zh"):
        """获取界面文字（中英文）"""
        texts_zh = {
            'app_title': 'JobsDB 自动求职工具',
            'tab_init': '1. 初始化配置',
            'tab_auto': '2. 全自动求职',
            'tab_manual': '3. 手动单岗处理',
            'tab_records': '4. 投递记录查询',
            'frame_api': 'API密钥配置',
            'frame_browser': '浏览器登录状态',
            'frame_user_info': '个人投递信息',
            'label_api_key': 'DeepSeek API Key：',
            'label_api_required': '必填',
            'label_chrome_profile': 'Chrome用户数据目录（可选）：',
            'label_chrome_profile_name': '配置文件名称：',
            'label_user_name': '姓名：',
            'label_user_email': '邮箱：',
            'label_user_phone': '电话：',
            'label_region': '地区选择：',
            'label_keyword': '搜索关键词：',
            'label_location': '搜索地点：',
            'label_threshold': '匹配度阈值：',
            'label_job_url': '岗位链接：',
            'label_job_description': '岗位描述：',
            'label_resume': '原始简历：',
            'label_resume_language': '简历语言：',
            'label_resume_language_auto': '自动检测',
            'button_open_website': '打开目标网站',
            'button_start_auto': '开始全自动求职',
            'button_pause': '暂停',
            'button_resume': '继续',
            'button_fetch_job': '抓取岗位信息',
            'button_generate': '生成定制简历',
            'button_export_pdf': '导出PDF',
            'button_export_records': '导出记录',
            'button_upload_word': '上传Word简历',
            'button_preview': '预览简历',
            'button_clear': '清空',
            'hint_auto_steps': '步骤：先完成「初始化配置」→ 上传简历 → 点击开始',
            'hint_init': '💡 提示：首次使用请先完成以下配置，所有信息会自动保存',
            'hint_browser_effect': '💡 作用：复用你Chrome里已登录的JobsDB账号（不用重复输密码）',
            'hint_browser_detail': '💡 不懂就这么填：路径框留空或填Chrome数据路径，配置文件名填Default即可。请确保浏览器中JobsDB账号已经登录状态。',
            'hint_auto_usage': '💡 提示：用于处理单个岗位，上传简历后生成定制简历，也可用于全自动求职前的简历准备',
            'hint_records': '💡 提示：查看所有已处理的岗位记录，包括匹配度和投递状态，可导出为Excel文件',
            'label_chrome_path': 'Chrome登录数据路径：',
            'label_chrome_profile_name': 'Chrome配置文件名：',
            'label_chrome_profile_hint': '（默认填Default，无需修改）',
            'frame_browser_title': '🌐 Chrome浏览器登录数据配置（可选）',
            'frame_api_title': '🔑 API密钥配置（已内置，无需填写）',
            'frame_user_title': '👤 个人投递信息（必填）',
            'frame_job_title': '📄 岗位信息',
            'frame_resume_title': '📝 原始简历（必填）',
            'frame_result_title': '✨ 生成的定制简历',
            'frame_auto_result_title': '📊 执行结果',
            'frame_records_title': '📋 投递记录列表',
            'label_required': '必填',
            'label_job_title': '岗位描述：',
            'label_resume_hint': '提示：上传Word简历或直接粘贴简历文本，此简历将用于生成定制简历',
            'label_resume_lang_hint': '（auto=自动检测，zh=中文，en=英文）',
            'button_browse': '浏览',
            'button_save': '💾 保存配置',
            'button_clear_cache': '🗑️ 清除缓存',
            'button_refresh': '🔄 刷新记录',
            'button_export': '📊 导出记录',
            'status_ready': '就绪',
            'warning': '警告',
            'error': '错误',
            'success': '成功',
            'confirm': '确认',
        }
        
        texts_en = {
            'app_title': 'JobsDB Auto Job Application Tool',
            'tab_init': '1. Initial Setup',
            'tab_auto': '2. Auto Job Search',
            'tab_manual': '3. Manual Single Job',
            'tab_records': '4. Application Records',
            'frame_api': 'API Key Configuration',
            'frame_browser': 'Browser Login Status',
            'frame_user_info': 'Personal Information',
            'label_api_key': 'DeepSeek API Key:',
            'label_api_required': 'Required',
            'label_chrome_profile': 'Chrome User Data Dir (Optional):',
            'label_chrome_profile_name': 'Profile Name:',
            'label_user_name': 'Name:',
            'label_user_email': 'Email:',
            'label_user_phone': 'Phone:',
            'label_region': 'Region:',
            'label_keyword': 'Search Keyword:',
            'label_location': 'Search Location:',
            'label_threshold': 'Match Threshold:',
            'label_job_url': 'Job Link:',
            'label_job_description': 'Job Description:',
            'label_resume': 'Original Resume:',
            'label_resume_language': 'Resume Language:',
            'label_resume_language_auto': 'Auto-detect',
            'button_open_website': 'Open Target Website',
            'button_start_auto': 'Start Auto Job Search',
            'button_pause': 'Pause',
            'button_resume': 'Resume',
            'button_fetch_job': 'Fetch Job Info',
            'button_generate': 'Generate Custom Resume',
            'button_export_pdf': 'Export PDF',
            'button_export_records': 'Export Records',
            'button_upload_word': 'Upload Word Resume',
            'button_preview': 'Preview Resume',
            'button_clear': 'Clear',
            'hint_auto_steps': 'Steps: Complete "Initial Setup" → Upload Resume → Click Start',
            'hint_init': '💡 Tip: Please complete the following configuration for first-time use. All information will be saved automatically.',
            'hint_browser_effect': '💡 Purpose: Reuse your logged-in JobsDB account in Chrome (no need to enter password repeatedly)',
            'hint_browser_detail': '💡 Don\'t know how? Leave the path field empty or fill in Chrome data path, profile name should be "Default". Please ensure your JobsDB account is logged in the browser.',
            'hint_auto_usage': '💡 Tip: Used to process a single job, upload resume to generate customized resume, or prepare resume before auto job search',
            'hint_records': '💡 Tip: View all processed job records, including match score and application status, can be exported to Excel',
            'label_chrome_path': 'Chrome Login Data Path:',
            'label_chrome_profile_name': 'Chrome Profile Name:',
            'label_chrome_profile_hint': '(Default: Default, no need to change)',
            'frame_browser_title': '🌐 Chrome Browser Login Data Configuration (Optional)',
            'frame_api_title': '🔑 API Key Configuration (Built-in, no need to fill)',
            'frame_user_title': '👤 Personal Information (Required)',
            'frame_job_title': '📄 Job Information',
            'frame_resume_title': '📝 Original Resume (Required)',
            'frame_result_title': '✨ Generated Customized Resume',
            'frame_auto_result_title': '📊 Execution Results',
            'frame_records_title': '📋 Application Records List',
            'label_required': 'Required',
            'label_job_title': 'Job Description:',
            'label_resume_hint': 'Tip: Upload Word resume or paste resume text directly. This resume will be used to generate customized resume',
            'label_resume_lang_hint': '(auto=auto-detect, zh=Chinese, en=English)',
            'button_browse': 'Browse',
            'button_save': '💾 Save Configuration',
            'button_clear_cache': '🗑️ Clear Cache',
            'button_refresh': '🔄 Refresh Records',
            'button_export': '📊 Export Records',
            'status_ready': 'Ready',
            'warning': 'Warning',
            'error': 'Error',
            'success': 'Success',
            'confirm': 'Confirm',
        }
        
        return texts_zh if lang == "zh" else texts_en
    
    def load_config(self):
        """加载配置文件"""
        # 首先加载API Key配置（开发者的API Key，不暴露给客户）
        # 支持打包后的exe路径
        api_config_file = get_resource_path("api_config.json")
        # 如果打包后的路径不存在，尝试当前目录
        if not os.path.exists(api_config_file):
            api_config_file = "api_config.json"
        
        default_api_key = ""
        use_proxy = False
        proxy_url = ""
        server_api_key = ""
        
        if os.path.exists(api_config_file):
            try:
                with open(api_config_file, 'r', encoding='utf-8') as f:
                    api_config = json.load(f)
                    default_api_key = api_config.get('api_key', '')
                    use_proxy = api_config.get('use_proxy', False)
                    proxy_url = api_config.get('proxy_url', 'http://localhost:5000')
                    server_api_key = api_config.get('server_api_key', '')
            except:
                pass
        
        # 加载用户配置
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.config = config
            except:
                self.config = {}
        else:
            self.config = {}
        
        # 如果用户配置中没有API Key，使用开发者的默认API Key
        if 'api_key' not in self.config or not self.config['api_key']:
            self.config['api_key'] = default_api_key
        
        # 设置代理配置
        self.config['use_proxy'] = use_proxy
        self.config['proxy_url'] = proxy_url
        self.config['server_api_key'] = server_api_key
    
    def save_config(self):
        """保存配置到文件"""
        try:
            # 保存配置时，不保存API Key（使用内置的）
            config_to_save = self.config.copy()
            if 'api_key' in config_to_save:
                del config_to_save['api_key']  # 不保存API Key到用户配置文件
            
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config_to_save, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置失败: {e}")
    
    def create_widgets(self):
        """创建主界面"""
        # 创建Notebook（标签页容器）
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建4个标签页
        self.tab_init = ttk.Frame(self.notebook, padding="10")
        self.tab_auto = ttk.Frame(self.notebook, padding="10")
        self.tab_manual = ttk.Frame(self.notebook, padding="10")
        self.tab_records = ttk.Frame(self.notebook, padding="10")
        
        self.notebook.add(self.tab_init, text=self.texts['tab_init'])
        self.notebook.add(self.tab_auto, text=self.texts['tab_auto'])
        self.notebook.add(self.tab_manual, text=self.texts['tab_manual'])
        self.notebook.add(self.tab_records, text=self.texts['tab_records'])
        
        # 创建各个标签页的内容
        self.create_tab_init()
        self.create_tab_auto()
        self.create_tab_manual()
        self.create_tab_records()
        
        # 创建顶部工具栏（右上角语言切换）
        top_toolbar = ttk.Frame(self.root)
        top_toolbar.pack(side=tk.TOP, fill=tk.X, padx=10, pady=5)
        
        # 语言切换按钮（右上角）
        lang_text = "EN" if self.language == "zh" else "中文"
        self.lang_button = ttk.Button(top_toolbar, text=lang_text, width=8, command=self.toggle_language)
        self.lang_button.pack(side=tk.RIGHT, padx=5)
        
        # 创建状态栏
        status_frame = ttk.Frame(self.root)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.status_label = ttk.Label(status_frame, text=self.texts['status_ready'], relief=tk.SUNKEN, anchor=tk.W)
        self.status_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 保存Chrome driver实例，以便复用
        self.chrome_driver = None
        self.chrome_debug_port = 9222  # 远程调试端口
    
    def create_tab_init(self):
        """创建标签1：初始化配置"""
        # 顶部提示
        hint_frame = ttk.Frame(self.tab_init)
        hint_frame.pack(fill=tk.X, pady=(0, 15))
        self.hint_init_label = ttk.Label(hint_frame, text=self.texts['hint_init'], 
                              foreground="blue", font=("Arial", 10, "bold"))
        self.hint_init_label.pack()
        
        # API密钥配置分组（隐藏，使用开发者的API Key）
        # 不显示给客户，API Key已内置
        self.api_frame = ttk.LabelFrame(self.tab_init, text=self.texts['frame_api_title'], padding="10")
        self.api_frame.pack(fill=tk.X, pady=(0, 10))
        
        api_key_row = ttk.Frame(self.api_frame)
        api_key_row.pack(fill=tk.X, pady=5)
        ttk.Label(api_key_row, text="API Key状态：", font=("Arial", 10)).pack(side=tk.LEFT)
        # 显示已配置状态，但不显示实际Key
        api_status = "✅ 已配置（内置）" if self.config.get('api_key') else "❌ 未配置"
        status_color = "green" if self.config.get('api_key') else "red"
        status_label = ttk.Label(api_key_row, text=api_status, 
                                foreground=status_color, font=("Arial", 10, "bold"))
        status_label.pack(side=tk.LEFT, padx=5)
        ttk.Label(api_key_row, text="（开发者已配置，客户无需填写）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=10)
        
        # 隐藏的API Key输入框（用于内部使用，不显示给客户）
        self.api_key_entry = ttk.Entry(api_key_row, width=1)
        self.api_key_entry.pack_forget()  # 隐藏输入框
        if 'api_key' in self.config:
            self.api_key_entry.insert(0, self.config['api_key'])
        
        # 浏览器登录状态分组
        self.browser_frame = ttk.LabelFrame(self.tab_init, text=self.texts['frame_browser_title'], padding="10")
        self.browser_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 自动检测Chrome路径
        chrome_path, chrome_detected, chrome_hint_msg = detect_chrome_user_data_dir()
        
        # 如果配置文件中已有路径，优先使用配置文件中的
        if 'chrome_user_data_dir' in self.config and self.config['chrome_user_data_dir']:
            chrome_path = self.config['chrome_user_data_dir']
            chrome_detected = True
            chrome_hint_msg = "使用已保存的配置" if self.language == "zh" else "Using saved configuration"
        
        # 说明文字
        self.browser_hint_label = ttk.Label(self.browser_frame, 
                                text=self.texts['hint_browser_effect'],
                                foreground="blue", font=("Arial", 9, "bold"))
        self.browser_hint_label.pack(anchor=tk.W, pady=(0, 5))
        
        # 详细说明
        self.browser_detail_label = ttk.Label(self.browser_frame, 
                                  text=self.texts['hint_browser_detail'],
                                  foreground="gray", font=("Arial", 9), wraplength=900)
        self.browser_detail_label.pack(anchor=tk.W, pady=(0, 8))
        
        # Chrome路径输入行
        chrome_dir_row = ttk.Frame(self.browser_frame)
        chrome_dir_row.pack(fill=tk.X, pady=5)
        self.chrome_path_label = ttk.Label(chrome_dir_row, text=self.texts['label_chrome_path'], font=("Arial", 10))
        self.chrome_path_label.pack(side=tk.LEFT)
        
        self.chrome_dir_entry = ttk.Entry(chrome_dir_row, width=50)
        self.chrome_dir_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 自动填充检测到的路径
        if chrome_path:
            self.chrome_dir_entry.insert(0, chrome_path)
            # 添加状态提示
            status_color = "green" if chrome_detected else "orange"
            self.chrome_status_label = ttk.Label(chrome_dir_row, text=f"✓ {chrome_hint_msg}", 
                                    foreground=status_color, font=("Arial", 9))
            self.chrome_status_label.pack(side=tk.LEFT, padx=5)
        else:
            # 未检测到Chrome，添加提示
            self.chrome_status_label = ttk.Label(chrome_dir_row, text="⚠ 未检测到Chrome" if self.language == "zh" else "⚠ Chrome not detected", 
                                    foreground="orange", font=("Arial", 9))
            self.chrome_status_label.pack(side=tk.LEFT, padx=5)
        
        self.browse_button = ttk.Button(chrome_dir_row, text=self.texts['button_browse'], command=self.browse_chrome_dir)
        self.browse_button.pack(side=tk.LEFT, padx=5)
        
        # 绑定路径变化事件，自动刷新配置文件列表
        def on_chrome_dir_change(event=None):
            chrome_path = self.chrome_dir_entry.get().strip()
            if chrome_path and os.path.exists(chrome_path):
                available_profiles = scan_chrome_profiles(chrome_path)
                if not available_profiles:
                    available_profiles = ["Default"]
                if hasattr(self, 'chrome_profile_combo'):
                    current_value = self.chrome_profile_var.get() if hasattr(self, 'chrome_profile_var') else "Default"
                    self.chrome_profile_combo['values'] = available_profiles
                    if current_value in available_profiles:
                        self.chrome_profile_var.set(current_value)
                    else:
                        self.chrome_profile_var.set(available_profiles[0])
        
        self.chrome_dir_entry.bind('<KeyRelease>', on_chrome_dir_change)
        self.chrome_dir_entry.bind('<FocusOut>', on_chrome_dir_change)
        
        # 路径提示（Windows示例）
        if sys.platform == "win32":
            example_path = "C:\\Users\\你的用户名\\AppData\\Local\\Google\\Chrome\\User Data"
        elif sys.platform == "darwin":
            example_path = "/Users/你的用户名/Library/Application Support/Google/Chrome"
        else:
            example_path = "~/.config/google-chrome"
        
        path_hint_text = f"💡 Windows默认路径示例：{example_path}" if self.language == "zh" else f"💡 Windows default path example: {example_path}"
        self.path_hint_label = ttk.Label(self.browser_frame, 
                             text=path_hint_text,
                             foreground="gray", font=("Arial", 8))
        self.path_hint_label.pack(anchor=tk.W, pady=(0, 5))
        
        # Chrome配置文件选择行（改为下拉框）
        chrome_profile_row = ttk.Frame(self.browser_frame)
        chrome_profile_row.pack(fill=tk.X, pady=5)
        self.chrome_profile_label = ttk.Label(chrome_profile_row, text="Chrome配置文件：" if self.language == "zh" else "Chrome Profile:", font=("Arial", 10))
        self.chrome_profile_label.pack(side=tk.LEFT)
        
        # 扫描可用的配置文件
        available_profiles = []
        if chrome_path:
            available_profiles = scan_chrome_profiles(chrome_path)
        
        # 如果没有扫描到，至少提供Default选项
        if not available_profiles:
            available_profiles = ["Default"]
        
        self.chrome_profile_var = tk.StringVar()
        self.chrome_profile_combo = ttk.Combobox(chrome_profile_row, textvariable=self.chrome_profile_var,
                                                 values=available_profiles, width=20, state="readonly")
        self.chrome_profile_combo.pack(side=tk.LEFT, padx=5)
        
        # 自动填充默认值
        if 'chrome_profile' in self.config and self.config['chrome_profile']:
            if self.config['chrome_profile'] in available_profiles:
                self.chrome_profile_var.set(self.config['chrome_profile'])
            else:
                self.chrome_profile_var.set(available_profiles[0])
        else:
            self.chrome_profile_var.set(available_profiles[0])
        
        # 刷新配置文件按钮
        refresh_profile_btn = ttk.Button(chrome_profile_row, text="🔄 刷新" if self.language == "zh" else "🔄 Refresh",
                                        command=self.refresh_chrome_profiles)
        refresh_profile_btn.pack(side=tk.LEFT, padx=5)
        
        self.chrome_profile_hint_label = ttk.Label(chrome_profile_row, 
                 text="（选择要使用的Chrome账号）" if self.language == "zh" else "(Select Chrome account to use)", 
                 foreground="gray", font=("Arial", 9))
        self.chrome_profile_hint_label.pack(side=tk.LEFT, padx=5)
        
        # 简历上传分组
        resume_frame_init = ttk.LabelFrame(self.tab_init, text="📝 上传简历（必填）" if self.language == "zh" else "📝 Upload Resume (Required)", padding="10")
        resume_frame_init.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        resume_hint_init = ttk.Label(resume_frame_init, 
                               text="提示：上传Word简历或直接粘贴简历文本，此简历将用于生成定制简历" if self.language == "zh" else "Tip: Upload Word resume or paste resume text directly. This resume will be used to generate customized resume",
                               foreground="gray", font=("Arial", 9))
        resume_hint_init.pack(anchor=tk.W, pady=(0, 8))
        
        # 简历文本区域和按钮
        resume_btn_row_init = ttk.Frame(resume_frame_init)
        resume_btn_row_init.pack(fill=tk.X, pady=(0, 5))
        upload_word_btn_init = ttk.Button(resume_btn_row_init, text="📤 上传Word简历" if self.language == "zh" else "📤 Upload Word Resume",
                                     command=self.on_upload_word_click_init)
        upload_word_btn_init.pack(side=tk.LEFT, padx=5)
        preview_btn_init = ttk.Button(resume_btn_row_init, text="👁️ 预览简历" if self.language == "zh" else "👁️ Preview Resume",
                                 command=self.on_preview_resume_click_init)
        preview_btn_init.pack(side=tk.LEFT, padx=5)
        
        self.resume_text_init = scrolledtext.ScrolledText(resume_frame_init, height=10, wrap=tk.WORD)
        self.resume_text_init.pack(fill=tk.BOTH, expand=True)
        # 加载简历内容（优先从配置，其次从缓存文件）
        if 'resume_content' in self.config and self.config['resume_content']:
            self.resume_text_init.insert("1.0", self.config['resume_content'])
        else:
            # 尝试从缓存文件加载
            cached_resume = self.load_resume_cache()
            if cached_resume:
                self.resume_text_init.insert("1.0", cached_resume)
                self.config['resume_content'] = cached_resume
        
        # 个人投递信息分组
        self.user_frame = ttk.LabelFrame(self.tab_init, text=self.texts['frame_user_title'], padding="10")
        self.user_frame.pack(fill=tk.X, pady=(0, 10))
        
        user_hint_text = "提示：这些信息将用于自动填写求职表单" if self.language == "zh" else "Tip: This information will be used to automatically fill in the job application form"
        user_hint = ttk.Label(self.user_frame, 
                             text=user_hint_text,
                             foreground="gray", font=("Arial", 9))
        user_hint.pack(anchor=tk.W, pady=(0, 8))
        
        name_row = ttk.Frame(self.user_frame)
        name_row.pack(fill=tk.X, pady=5)
        name_label_text = "姓名：" if self.language == "zh" else "Name:"
        ttk.Label(name_row, text=name_label_text, font=("Arial", 10)).pack(side=tk.LEFT)
        ttk.Label(name_row, text=self.texts['label_required'], foreground="red", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        self.user_name_entry = ttk.Entry(name_row, width=30)
        self.user_name_entry.pack(side=tk.LEFT, padx=5)
        if 'user_name' in self.config:
            self.user_name_entry.insert(0, self.config['user_name'])
        
        email_row = ttk.Frame(self.user_frame)
        email_row.pack(fill=tk.X, pady=5)
        email_label_text = "邮箱：" if self.language == "zh" else "Email:"
        ttk.Label(email_row, text=email_label_text, font=("Arial", 10)).pack(side=tk.LEFT)
        ttk.Label(email_row, text=self.texts['label_required'], foreground="red", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        self.user_email_entry = ttk.Entry(email_row, width=30)
        self.user_email_entry.pack(side=tk.LEFT, padx=5)
        if 'user_email' in self.config:
            self.user_email_entry.insert(0, self.config['user_email'])
        
        phone_row = ttk.Frame(self.user_frame)
        phone_row.pack(fill=tk.X, pady=5)
        phone_label_text = "电话：" if self.language == "zh" else "Phone:"
        ttk.Label(phone_row, text=phone_label_text, font=("Arial", 10)).pack(side=tk.LEFT)
        ttk.Label(phone_row, text=self.texts['label_required'], foreground="red", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        self.user_phone_entry = ttk.Entry(phone_row, width=30)
        self.user_phone_entry.pack(side=tk.LEFT, padx=5)
        if 'user_phone' in self.config:
            self.user_phone_entry.insert(0, self.config['user_phone'])
        
        # 期望薪资
        salary_row = ttk.Frame(self.user_frame)
        salary_row.pack(fill=tk.X, pady=5)
        salary_label_text = "期望薪资：" if self.language == "zh" else "Expected Salary:"
        ttk.Label(salary_row, text=salary_label_text, font=("Arial", 10)).pack(side=tk.LEFT)
        ttk.Label(salary_row, text=self.texts['label_required'], foreground="red", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        self.expected_salary_var = tk.StringVar()
        salary_values = [
            "$5K", "$6K", "$7K", "$8K", "$9K", "$10K", "$11K", "$12K", "$13K", "$14K", 
            "$15K", "$16K", "$17K", "$18K", "$19K", "$20K", "$25K", "$30K", "$35K", 
            "$40K", "$45K", "$50K", "$60K", "$70K", "$80K", "$90K", "$100K", "$120K or more"
        ]
        salary_combo = ttk.Combobox(salary_row, textvariable=self.expected_salary_var,
                                    values=salary_values, width=20, state="readonly")
        salary_combo.pack(side=tk.LEFT, padx=5)
        if 'expected_salary' in self.config:
            self.expected_salary_var.set(self.config['expected_salary'])
        else:
            self.expected_salary_var.set("$20K")
        salary_hint_text = "（用于自动填写JobsDB申请表单）" if self.language == "zh" else "(For auto-filling JobsDB application form)"
        ttk.Label(salary_row, text=salary_hint_text, 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 保存配置按钮和清除缓存按钮
        button_row = ttk.Frame(self.tab_init)
        button_row.pack(pady=15)
        
        self.save_button = tk.Button(button_row, text=self.texts['button_save'], 
                               command=self.save_config_from_ui,
                               bg="#4A90E2", fg="white", font=("Arial", 10, "bold"),
                               padx=15, pady=5, cursor="hand2")
        self.save_button.pack(side=tk.LEFT, padx=5)
        
        self.clear_cache_btn = ttk.Button(button_row, text=self.texts['button_clear_cache'], command=self.clear_cache)
        self.clear_cache_btn.pack(side=tk.LEFT, padx=5)
    
    def refresh_chrome_profiles(self):
        """刷新Chrome配置文件列表"""
        chrome_path = self.chrome_dir_entry.get().strip()
        if chrome_path and os.path.exists(chrome_path):
            available_profiles = scan_chrome_profiles(chrome_path)
            if not available_profiles:
                available_profiles = ["Default"]
            
            # 更新下拉框选项
            if hasattr(self, 'chrome_profile_combo'):
                current_value = self.chrome_profile_var.get() if hasattr(self, 'chrome_profile_var') else "Default"
                self.chrome_profile_combo['values'] = available_profiles
                # 如果当前值还在列表中，保持；否则选择第一个
                if current_value in available_profiles:
                    self.chrome_profile_var.set(current_value)
                else:
                    self.chrome_profile_var.set(available_profiles[0])
                messagebox.showinfo("成功", f"已刷新配置文件列表，找到 {len(available_profiles)} 个配置文件" if self.language == "zh" else f"Refreshed profile list, found {len(available_profiles)} profiles")
            else:
                messagebox.showwarning("警告", "配置文件下拉框未初始化" if self.language == "zh" else "Profile combo box not initialized")
        else:
            messagebox.showwarning("警告", "请先设置Chrome用户数据目录路径" if self.language == "zh" else "Please set Chrome user data directory first")
    
    def browse_chrome_dir(self):
        """浏览Chrome用户数据目录"""
        directory = filedialog.askdirectory(title="选择Chrome用户数据目录")
        if directory:
            self.chrome_dir_entry.delete(0, tk.END)
            self.chrome_dir_entry.insert(0, directory)
    
    def save_config_from_ui(self):
        """从UI保存配置"""
        # 不保存API Key到用户配置，使用内置的API Key
        # self.config['api_key'] = self.api_key_entry.get()  # 注释掉，使用内置API Key
        self.config['chrome_user_data_dir'] = self.chrome_dir_entry.get()
        if hasattr(self, 'chrome_profile_var'):
            self.config['chrome_profile'] = self.chrome_profile_var.get()
        elif hasattr(self, 'chrome_profile_entry'):
            self.config['chrome_profile'] = self.chrome_profile_entry.get()
        self.config['user_name'] = self.user_name_entry.get()
        self.config['user_email'] = self.user_email_entry.get()
        self.config['user_phone'] = self.user_phone_entry.get()
        self.save_config()
        messagebox.showinfo(self.texts['success'], "配置已保存！")
    
    def clear_cache(self):
        """清除缓存"""
        if messagebox.askyesno("确认", "确定要清除所有缓存吗？\n\n这将清除：\n- 简历缓存\n- 配置信息\n\n此操作不可恢复！"):
            try:
                # 清除简历缓存文件
                cache_file = "resume_cache.txt"
                if os.path.exists(cache_file):
                    os.remove(cache_file)
                
                # 清除配置文件
                if os.path.exists(self.config_file):
                    os.remove(self.config_file)
                
                # 重置配置
                self.config = {}
                
                # 清空UI
                if hasattr(self, 'api_key_entry'):
                    self.api_key_entry.delete(0, tk.END)
                if hasattr(self, 'chrome_dir_entry'):
                    self.chrome_dir_entry.delete(0, tk.END)
                if hasattr(self, 'chrome_profile_entry'):
                    self.chrome_profile_entry.delete(0, tk.END)
                if hasattr(self, 'user_name_entry'):
                    self.user_name_entry.delete(0, tk.END)
                if hasattr(self, 'user_email_entry'):
                    self.user_email_entry.delete(0, tk.END)
                if hasattr(self, 'user_phone_entry'):
                    self.user_phone_entry.delete(0, tk.END)
                if hasattr(self, 'resume_text'):
                    self.resume_text.delete("1.0", tk.END)
                
                messagebox.showinfo(self.texts['success'], "缓存已清除！")
            except Exception as e:
                messagebox.showerror(self.texts['error'], f"清除缓存失败: {str(e)}")
    
    def create_tab_auto(self):
        """创建标签2：全自动求职"""
        # 顶部流程指引
        hint_frame = ttk.LabelFrame(self.tab_auto, text="📋 使用流程" if self.language == "zh" else "📋 Usage Steps", padding="10")
        hint_frame.pack(fill=tk.X, pady=(0, 15))
        # 更新提示语，因为简历现在在标签1
        hint_auto_text = "步骤1：完成「初始化配置」标签页的配置（包括上传简历） → 步骤2：填写搜索条件 → 步骤3：点击「开始全自动求职」"
        self.hint_auto_label = ttk.Label(hint_frame, text=hint_auto_text, 
                              foreground="blue", font=("Arial", 10, "bold"), wraplength=900)
        self.hint_auto_label.pack()
        
        # 搜索条件分组
        self.search_frame = ttk.LabelFrame(self.tab_auto, text="🔍 搜索条件" if self.language == "zh" else "🔍 Search Criteria", padding="10")
        self.search_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 地区选择
        region_row = ttk.Frame(self.search_frame)
        region_row.pack(fill=tk.X, pady=5)
        ttk.Label(region_row, text="地区选择：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.region_var = tk.StringVar()
        region_combo = ttk.Combobox(region_row, textvariable=self.region_var, 
                                    values=["香港 (hk)", "新加坡 (sg)", "马来西亚 (my)", "菲律宾 (ph)"],
                                    width=20, state="readonly")
        region_combo.pack(side=tk.LEFT, padx=5)
        if 'region' in self.config:
            self.region_var.set(self.config['region'])
        else:
            self.region_var.set("香港 (hk)")
        
        # 搜索关键词
        keyword_row = ttk.Frame(self.search_frame)
        keyword_row.pack(fill=tk.X, pady=5)
        ttk.Label(keyword_row, text="搜索关键词：", font=("Arial", 10)).pack(side=tk.LEFT)
        ttk.Label(keyword_row, text="必填", foreground="red", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        self.search_keyword_entry = ttk.Entry(keyword_row, width=40)
        self.search_keyword_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        if 'search_keyword' in self.config:
            self.search_keyword_entry.insert(0, self.config['search_keyword'])
        ttk.Label(keyword_row, text="（例如：Administrative Officer）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 行业分类
        category_row = ttk.Frame(self.search_frame)
        category_row.pack(fill=tk.X, pady=5)
        ttk.Label(category_row, text="行业分类：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.category_var = tk.StringVar()
        category_values = [
            "全部",
            "Accounting",
            "Administration & Office Support",
            "Advertising, Arts & Media",
            "Banking & Financial Services",
            "Call Centre & Customer Service",
            "CEO & General Management",
            "Community Services & Development",
            "Construction",
            "Consulting & Strategy",
            "Design & Architecture",
            "Education & Training",
            "Engineering",
            "Farming, Animals & Conservation",
            "Government & Defence",
            "Healthcare & Medical",
            "Hospitality & Tourism",
            "Human Resources & Recruitment",
            "Information & Communication Technology",
            "Insurance & Superannuation",
            "Legal",
            "Manufacturing, Transport & Logistics",
            "Marketing & Communications",
            "Mining, Resources & Energy",
            "Real Estate & Property",
            "Retail & Consumer Products",
            "Sales",
            "Science & Technology",
            "Self Employment",
            "Sport & Recreation",
            "Trades & Services"
        ]
        category_combo = ttk.Combobox(category_row, textvariable=self.category_var,
                                      values=category_values, width=45, state="readonly")
        category_combo.pack(side=tk.LEFT, padx=5)
        if 'job_category' in self.config:
            self.category_var.set(self.config['job_category'])
        else:
            self.category_var.set("全部")
        ttk.Label(category_row, text="（可选）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 搜索地点
        location_row = ttk.Frame(self.search_frame)
        location_row.pack(fill=tk.X, pady=5)
        ttk.Label(location_row, text="搜索地点：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.search_location_entry = ttk.Entry(location_row, width=40)
        self.search_location_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        if 'search_location' in self.config:
            self.search_location_entry.insert(0, self.config['search_location'])
        ttk.Label(location_row, text="（例如：Hong Kong，可选）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 匹配度阈值
        threshold_row = ttk.Frame(self.search_frame)
        threshold_row.pack(fill=tk.X, pady=5)
        ttk.Label(threshold_row, text="匹配度阈值：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.match_threshold_entry = ttk.Entry(threshold_row, width=10)
        self.match_threshold_entry.pack(side=tk.LEFT, padx=5)
        if 'match_threshold' in self.config:
            self.match_threshold_entry.insert(0, str(self.config['match_threshold']))
        else:
            self.match_threshold_entry.insert(0, "70")
        ttk.Label(threshold_row, text="（0-100，建议70，只投递匹配度≥此值的岗位）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 投递控制分组
        apply_control_frame = ttk.LabelFrame(self.tab_auto, text="⚙️ 投递控制（安全设置）" if self.language == "zh" else "⚙️ Application Control (Safety Settings)", padding="10")
        apply_control_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 投递总数量
        max_apply_row = ttk.Frame(apply_control_frame)
        max_apply_row.pack(fill=tk.X, pady=5)
        ttk.Label(max_apply_row, text="投递总数量：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.max_apply_count_entry = ttk.Entry(max_apply_row, width=10)
        self.max_apply_count_entry.pack(side=tk.LEFT, padx=5)
        if 'max_apply_count' in self.config:
            self.max_apply_count_entry.insert(0, str(self.config['max_apply_count']))
        else:
            self.max_apply_count_entry.insert(0, "15")
        ttk.Label(max_apply_row, text="（最大15个，建议≤15，避免被检测）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 投递间隔（最小）
        interval_min_row = ttk.Frame(apply_control_frame)
        interval_min_row.pack(fill=tk.X, pady=5)
        ttk.Label(interval_min_row, text="投递间隔（最小）：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.apply_interval_min_entry = ttk.Entry(interval_min_row, width=10)
        self.apply_interval_min_entry.pack(side=tk.LEFT, padx=5)
        if 'apply_interval_min' in self.config:
            self.apply_interval_min_entry.insert(0, str(self.config['apply_interval_min']))
        else:
            self.apply_interval_min_entry.insert(0, "6")
        ttk.Label(interval_min_row, text="分钟", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=2)
        
        # 投递间隔（最大）
        interval_max_row = ttk.Frame(apply_control_frame)
        interval_max_row.pack(fill=tk.X, pady=5)
        ttk.Label(interval_max_row, text="投递间隔（最大）：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.apply_interval_max_entry = ttk.Entry(interval_max_row, width=10)
        self.apply_interval_max_entry.pack(side=tk.LEFT, padx=5)
        if 'apply_interval_max' in self.config:
            self.apply_interval_max_entry.insert(0, str(self.config['apply_interval_max']))
        else:
            self.apply_interval_max_entry.insert(0, "12")
        ttk.Label(interval_max_row, text="分钟（系统会在最小和最大之间随机选择，模拟真人操作）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        
        # 安全提示
        safety_hint = ttk.Label(apply_control_frame, 
                               text="💡 安全提示：单日投递上限为15个，系统会自动检查并限制。投递间隔建议6-12分钟，避免被JobsDB判定为机器人。",
                               foreground="blue", font=("Arial", 9), wraplength=900)
        safety_hint.pack(anchor=tk.W, pady=(5, 0))
        
        # 按钮区域
        button_frame = ttk.Frame(self.tab_auto)
        button_frame.pack(pady=20)
        
        # 打开目标网站按钮
        open_website_btn = ttk.Button(button_frame, text="🌐 打开目标网站",
                                      command=self.on_open_website_click)
        open_website_btn.pack(side=tk.LEFT, padx=5)
        
        # 开始全自动求职按钮（高亮）
        self.start_auto_btn = tk.Button(button_frame, text="🚀 开始全自动求职",
                                        command=self.on_start_auto_click,
                                        bg="#4A90E2", fg="white", font=("Arial", 12, "bold"),
                                        padx=25, pady=8, cursor="hand2")
        self.start_auto_btn.pack(side=tk.LEFT, padx=10)
        
        # 暂停/继续按钮
        self.pause_button = ttk.Button(button_frame, text="⏸️ 暂停",
                                       command=self.on_pause_click, state="disabled")
        self.pause_button.pack(side=tk.LEFT, padx=5)
        
        # 结果显示区域
        self.result_frame = ttk.LabelFrame(self.tab_auto, text=self.texts['frame_auto_result_title'], padding="10")
        self.result_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        self.auto_result_text = scrolledtext.ScrolledText(self.result_frame, height=15, wrap=tk.WORD)
        self.auto_result_text.pack(fill=tk.BOTH, expand=True)
    
    def on_open_website_click(self):
        """打开目标网站（使用Chrome远程调试，以便后续复用）"""
        region = self.config.get('region', '香港 (hk)')
        if 'hk' in region:
            url = 'https://hk.jobsdb.com'
        elif 'sg' in region:
            url = 'https://sg.jobsdb.com'
        elif 'my' in region:
            url = 'https://my.jobsdb.com'
        elif 'ph' in region:
            url = 'https://ph.jobsdb.com'
        else:
            url = 'https://hk.jobsdb.com'
        
        # 使用统一的get_chrome_driver函数（带进程检查）
        driver, error_or_warning = self.get_chrome_driver(check_running=True)
        if driver is None:
            # 这是真正的错误
            messagebox.showerror("错误", error_or_warning)
            self.update_status("打开目标网站失败")
            return
        elif error_or_warning:
            # 这是警告（成功启动但未使用用户数据目录），记录但继续执行
            print(f"警告: {error_or_warning}")
        
        try:
            driver.get(url)
            self.update_status("已打开目标网站（已启用远程调试，后续操作将复用此浏览器）")
        except Exception as e:
            messagebox.showerror("错误", f"打开网站失败: {str(e)}")
            self.update_status("打开目标网站失败")
    
    def on_start_auto_click(self):
        """开始全自动求职"""
        if self.is_auto_running:
            # 停止
            self.is_auto_running = False
            self.start_auto_btn.config(text=self.texts['button_start_auto'])
            self.pause_button.config(state="disabled")
            self.update_status("已停止自动求职")
            return
        
        # 检查必要配置
        if not self.config.get('api_key'):
            messagebox.showerror(self.texts['error'], "请先配置API Key")
            return
        
        # 检查简历内容（从标签1的简历区域或配置中读取）
        resume_content = ""
        if hasattr(self, 'resume_text_init'):
            resume_content = self.resume_text_init.get("1.0", tk.END).strip()
        else:
            resume_content = self.config.get('resume_content', '').strip()
        
        if not resume_content:
            messagebox.showerror(self.texts['error'], "请先上传或输入原始简历")
            return
        
        # 检查Chrome配置（如果配置了用户数据目录，确保路径有效）
        user_data_dir = self.config.get('chrome_user_data_dir', '').strip()
        if user_data_dir and not os.path.exists(user_data_dir):
            messagebox.showerror(self.texts['error'], "Chrome用户数据目录路径不存在，请在「初始化配置」中检查路径设置")
            return
        
        # 开始自动求职
        self.is_auto_running = True
        self.start_auto_btn.config(text="停止自动求职" if self.language == "zh" else "Stop Auto Search")
        self.pause_button.config(state="normal")
        
        # 在新线程中运行
        thread = threading.Thread(target=self.auto_job_search_worker, daemon=True)
        thread.start()
    
    def auto_job_search_worker(self):
        """自动求职工作线程"""
        try:
            keyword = self.search_keyword_entry.get().strip()
            location = self.search_location_entry.get().strip()
            threshold = int(self.match_threshold_entry.get() or 70)
            region = self.region_var.get()
            
            # 读取投递控制设置
            max_apply_count = self.config.get('max_apply_count', 15)
            if hasattr(self, 'max_apply_count_entry'):
                try:
                    max_apply_count = int(self.max_apply_count_entry.get())
                    if max_apply_count > 15:
                        max_apply_count = 15
                except:
                    max_apply_count = 15
            
            apply_interval_min = self.config.get('apply_interval_min', 6)
            if hasattr(self, 'apply_interval_min_entry'):
                try:
                    apply_interval_min = int(self.apply_interval_min_entry.get())
                    if apply_interval_min < 5:
                        apply_interval_min = 5
                except:
                    apply_interval_min = 6
            
            apply_interval_max = self.config.get('apply_interval_max', 12)
            if hasattr(self, 'apply_interval_max_entry'):
                try:
                    apply_interval_max = int(self.apply_interval_max_entry.get())
                    if apply_interval_max < 10:
                        apply_interval_max = 10
                except:
                    apply_interval_max = 12
            
            # 检查单日投递数量
            daily_count = self.get_daily_apply_count()
            remaining_daily_quota = 15 - daily_count
            if remaining_daily_quota <= 0:
                self.log_auto_result(f"⚠️ 今日已投递15个岗位，已达到单日上限，请明天再试\n")
                return
            
            # 从标签1的简历区域或配置中读取简历
            if hasattr(self, 'resume_text_init'):
                original_resume = self.resume_text_init.get("1.0", tk.END).strip()
            else:
                original_resume = self.config.get('resume_content', '').strip()
            
            if not keyword:
                self.log_auto_result("错误：搜索关键词不能为空\n")
                return
            
            if not original_resume:
                self.log_auto_result("错误：原始简历不能为空\n")
                return
            
            api_key = self.config.get('api_key', '')
            if not api_key and hasattr(self, 'api_key_entry'):
                api_key = self.api_key_entry.get().strip()
            
            if not api_key:
                self.log_auto_result("错误：API Key未配置\n")
                return
            
            self.log_auto_result(f"开始搜索：关键词={keyword}, 地点={location}, 阈值={threshold}%\n")
            self.log_auto_result(f"投递控制：最大投递数={max_apply_count}, 今日剩余配额={remaining_daily_quota}, 投递间隔={apply_interval_min}-{apply_interval_max}分钟\n\n")
            
            # 步骤1：搜索并抓取岗位URL
            self.log_auto_result("正在搜索岗位...\n")
            search_criteria = {'keyword': keyword, 'location': location}
            success, result = self.scrape_job_urls(search_criteria, max_pages=3)
            
            if not success:
                self.log_auto_result(f"搜索失败：{result}\n")
                return
            
            job_urls = result
            if not job_urls:
                self.log_auto_result("未找到任何岗位\n")
                return
            
            self.log_auto_result(f"找到 {len(job_urls)} 个岗位，开始筛选...\n\n")
            
            # 步骤2：对每个岗位计算匹配度并筛选
            matched_jobs = []
            processed = 0
            
            for i, job_url in enumerate(job_urls, 1):
                if not self.is_auto_running:
                    self.log_auto_result("已停止\n")
                    break
                
                # 等待暂停事件
                self.pause_event.wait()
                
                self.log_auto_result(f"处理第 {i}/{len(job_urls)} 个岗位...\n")
                
                try:
                    # 抓取岗位描述
                    job_info, error = self.fetch_job_info(job_url)
                    if error:
                        self.log_auto_result(f"  抓取失败：{error}\n")
                        continue
                    
                    job_description = job_info.get('description', '')
                    if not job_description:
                        self.log_auto_result(f"  岗位描述为空，跳过\n")
                        continue
                    
                    # 计算匹配度
                    match_score = self.calculate_match_score(job_description, original_resume)
                    self.log_auto_result(f"  匹配度：{match_score}%\n")
                    
                    # 筛选：只保留匹配度>=阈值的岗位
                    if match_score >= threshold:
                        matched_jobs.append({
                            'url': job_url,
                            'title': job_info.get('title', 'Unknown'),
                            'description': job_description,
                            'match_score': match_score
                        })
                        self.log_auto_result(f"  ✅ 匹配度达标，已加入队列\n")
                    else:
                        self.log_auto_result(f"  ❌ 匹配度不足，已跳过\n")
                    
                    processed += 1
                    
                    # 随机延时
                    if i < len(job_urls):
                        delay = random.randint(3, 6)
                        time.sleep(delay)
                    
                except Exception as e:
                    self.log_auto_result(f"  处理异常：{str(e)}\n")
                    continue
            
            # 步骤3：生成定制简历并投递（受投递控制限制）
            if matched_jobs:
                self.log_auto_result(f"\n找到 {len(matched_jobs)} 个匹配岗位，开始生成简历并投递...\n\n")
                
                applied_count = 0  # 本次已投递数量
                first_apply = True  # 标记是否为第一次投递（第一次不需要等待）
                
                for i, job in enumerate(matched_jobs, 1):
                    if not self.is_auto_running:
                        break
                    
                    # 检查是否达到最大投递数量
                    if applied_count >= max_apply_count:
                        self.log_auto_result(f"\n⚠️ 已达到本次最大投递数量限制（{max_apply_count}个），停止投递\n")
                        break
                    
                    # 检查单日投递上限
                    current_daily_count = self.get_daily_apply_count()
                    if current_daily_count >= 15:
                        self.log_auto_result(f"\n⚠️ 今日已投递15个岗位，已达到单日上限，停止投递\n")
                        break
                    
                    self.pause_event.wait()
                    
                    self.log_auto_result(f"处理岗位 {i}/{len(matched_jobs)}：{job['title']}\n")
                    self.log_auto_result(f"  今日已投递：{current_daily_count}/15，本次已投递：{applied_count}/{max_apply_count}\n")
                    
                    try:
                        # 生成定制简历
                        self.log_auto_result(f"  正在生成定制简历...\n")
                        custom_resume, error = self.generate_custom_resume(
                            job['description'], original_resume, "auto"
                        )
                        
                        if error:
                            self.log_auto_result(f"  简历生成失败：{error}\n")
                            continue
                        
                        # 生成Cover Letter
                        self.log_auto_result(f"  正在生成Cover Letter...\n")
                        cover_letter, error = self.generate_cover_letter(
                            job['description'], 
                            job['title'], 
                            "Unknown",  # 公司名称，可以从job_info中获取
                            original_resume
                        )
                        
                        if error:
                            self.log_auto_result(f"  Cover Letter生成失败：{error}，将使用默认文本\n")
                            cover_letter = f"Dear Hiring Manager,\n\nI am writing to apply for the {job['title']} position. I believe my skills and experience make me a strong candidate for this role.\n\nSincerely,\n{self.config.get('user_name', '')}"
                        
                        # 将简历转换为PDF
                        self.log_auto_result(f"  正在转换简历为PDF...\n")
                        resume_pdf_path = f"resume_{int(time.time())}.pdf"
                        pdf_path, error = self.convert_resume_to_pdf(custom_resume, resume_pdf_path)
                        
                        if error:
                            self.log_auto_result(f"  PDF转换失败：{error}，将尝试直接上传文本\n")
                            resume_pdf_path = None
                        
                        # 准备用户信息
                        user_info = {
                            'name': self.config.get('user_name', ''),
                            'email': self.config.get('user_email', ''),
                            'phone': self.config.get('user_phone', ''),
                            'expected_salary': self.config.get('expected_salary', '$20K')
                        }
                        
                        # 自动投递
                        self.log_auto_result(f"  正在自动投递...\n")
                        # 构建申请URL（从岗位详情页跳转到申请页）
                        apply_url = job['url']
                        # JobsDB的申请URL通常是原URL加上/apply/或直接访问申请页面
                        # 先尝试访问岗位详情页，然后点击申请按钮
                        if '/job/' in apply_url and '/apply/' not in apply_url:
                            # 尝试构建申请URL
                            apply_url = apply_url.replace('/job/', '/apply/')
                        # 如果URL已经是申请页，直接使用
                        
                        success, message = self.auto_apply_job(
                            apply_url,
                            custom_resume,
                            cover_letter,
                            user_info,
                            resume_pdf_path
                        )
                        
                        if success:
                            # 保存记录（标记为已投递）
                            self.save_application_record(
                                job['title'], 
                                "Unknown", 
                                job['url'], 
                                job['match_score'], 
                                "已投递"
                            )
                            applied_count += 1
                            self.log_auto_result(f"  ✅ 投递成功（第{applied_count}个）\n")
                        else:
                            self.log_auto_result(f"  ❌ 投递失败：{message}\n")
                        
                        # 清理临时PDF文件
                        if resume_pdf_path and os.path.exists(resume_pdf_path):
                            try:
                                os.remove(resume_pdf_path)
                            except:
                                pass
                        
                        # 投递间隔控制（除了第一次投递）
                        if not first_apply and applied_count < max_apply_count:
                            # 随机选择间隔时间（分钟转秒）
                            interval_seconds = random.randint(apply_interval_min, apply_interval_max) * 60
                            interval_minutes = interval_seconds / 60
                            self.log_auto_result(f"  ⏳ 等待 {interval_minutes:.1f} 分钟后继续投递（模拟真人操作）...\n")
                            
                            # 分段等待，以便可以响应暂停/停止
                            wait_chunks = interval_seconds // 10  # 每10秒检查一次
                            for _ in range(wait_chunks):
                                if not self.is_auto_running:
                                    break
                                self.pause_event.wait()
                                time.sleep(10)
                            
                            # 剩余时间
                            remaining_seconds = interval_seconds % 10
                            if remaining_seconds > 0 and self.is_auto_running:
                                self.pause_event.wait()
                                time.sleep(remaining_seconds)
                        else:
                            first_apply = False
                        
                    except Exception as e:
                        self.log_auto_result(f"  处理异常：{str(e)}\n")
                        continue
                
                self.log_auto_result(f"\n✅ 完成！共处理 {processed} 个岗位，匹配 {len(matched_jobs)} 个，本次投递 {applied_count} 个\n")
                final_daily_count = self.get_daily_apply_count()
                self.log_auto_result(f"📊 今日累计投递：{final_daily_count}/15\n")
            else:
                self.log_auto_result(f"\n未找到匹配度>= {threshold}% 的岗位\n")
            
        except Exception as e:
            self.log_auto_result(f"错误: {str(e)}\n")
        finally:
            self.is_auto_running = False
            self.root.after(0, lambda: self.start_auto_btn.config(text=self.texts['button_start_auto']))
            self.root.after(0, lambda: self.pause_button.config(state="disabled"))
    
    def log_auto_result(self, message):
        """在自动求职结果区域添加日志"""
        self.root.after(0, lambda: self.auto_result_text.insert(tk.END, message))
        self.root.after(0, lambda: self.auto_result_text.see(tk.END))
    
    def on_pause_click(self):
        """暂停/继续"""
        if self.is_paused:
            self.is_paused = False
            self.pause_event.set()
            self.pause_button.config(text=self.texts['button_pause'])
            self.update_status("已继续")
        else:
            self.is_paused = True
            self.pause_event.clear()
            self.pause_button.config(text=self.texts['button_resume'])
            self.update_status("已暂停")
    
    def update_status(self, message):
        """更新状态栏"""
        self.root.after(0, lambda: self.status_label.config(text=message))
    
    def create_tab_manual(self):
        """创建标签3：手动单岗处理"""
        # 顶部提示
        hint_frame = ttk.Frame(self.tab_manual)
        hint_frame.pack(fill=tk.X, pady=(0, 15))
        hint_text = "💡 提示：用于处理单个岗位，生成定制简历。简历已在「初始化配置」标签页上传"
        hint_label = ttk.Label(hint_frame, text=hint_text, 
                              foreground="blue", font=("Arial", 10, "bold"), wraplength=900)
        hint_label.pack()
        
        # 岗位信息分组
        job_frame = ttk.LabelFrame(self.tab_manual, text="📄 岗位信息", padding="10")
        job_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # 岗位链接
        url_row = ttk.Frame(job_frame)
        url_row.pack(fill=tk.X, pady=5)
        ttk.Label(url_row, text="岗位链接：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.job_url_entry = ttk.Entry(url_row, width=60)
        self.job_url_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        fetch_btn = ttk.Button(url_row, text="🔍 抓取岗位信息", 
                               command=self.on_fetch_job_click)
        fetch_btn.pack(side=tk.LEFT, padx=5)
        
        # 岗位描述
        desc_label = ttk.Label(job_frame, text="岗位描述：", font=("Arial", 10))
        desc_label.pack(anchor=tk.W, pady=(10, 5))
        self.job_description_text = scrolledtext.ScrolledText(job_frame, height=12, wrap=tk.WORD)
        self.job_description_text.pack(fill=tk.BOTH, expand=True)
        
        # 简历语言选择（用于生成定制简历）
        lang_frame = ttk.Frame(self.tab_manual)
        lang_frame.pack(fill=tk.X, pady=(10, 5))
        ttk.Label(lang_frame, text="简历语言：", font=("Arial", 10)).pack(side=tk.LEFT)
        self.resume_language_var = tk.StringVar(value="auto")
        lang_combo = ttk.Combobox(lang_frame, textvariable=self.resume_language_var,
                                  values=["auto", "zh", "en"], width=15, state="readonly")
        lang_combo.pack(side=tk.LEFT, padx=5)
        ttk.Label(lang_frame, text="（auto=自动检测，zh=中文，en=英文）", 
                 foreground="gray", font=("Arial", 9)).pack(side=tk.LEFT, padx=5)
        ttk.Label(lang_frame, text="提示：简历已在「初始化配置」标签页上传", 
                 foreground="blue", font=("Arial", 9)).pack(side=tk.LEFT, padx=10)
        
        # 操作按钮区域
        action_frame = ttk.Frame(self.tab_manual)
        action_frame.pack(fill=tk.X, pady=10)
        
        generate_btn = tk.Button(action_frame, text="✨ 生成定制简历",
                                command=self.on_generate_click,
                                bg="#4A90E2", fg="white", font=("Arial", 11, "bold"),
                                padx=20, pady=5, cursor="hand2")
        generate_btn.pack(side=tk.LEFT, padx=5)
        
        export_pdf_btn = ttk.Button(action_frame, text="📄 导出PDF",
                                    command=self.on_export_pdf_click)
        export_pdf_btn.pack(side=tk.LEFT, padx=5)
        
        clear_btn = ttk.Button(action_frame, text="🗑️ 清空",
                              command=self.on_clear_click)
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # 生成结果区域
        result_frame = ttk.LabelFrame(self.tab_manual, text="✨ 生成的定制简历", padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        self.result_text = scrolledtext.ScrolledText(result_frame, height=12, wrap=tk.WORD)
        self.result_text.pack(fill=tk.BOTH, expand=True)
    
    def on_fetch_job_click(self):
        """抓取岗位信息"""
        job_url = self.job_url_entry.get().strip()
        if not job_url:
            messagebox.showerror(self.texts['error'], "请输入岗位链接")
            return
        
        self.update_status("正在抓取岗位信息...")
        
        def fetch_worker():
            try:
                job_info, error = self.fetch_job_info(job_url)
                if error:
                    self.root.after(0, lambda: messagebox.showerror(self.texts['error'], error))
                    self.root.after(0, lambda: self.update_status(self.texts['status_ready']))
                else:
                    # 更新UI
                    self.root.after(0, lambda: self.job_description_text.delete("1.0", tk.END))
                    self.root.after(0, lambda: self.job_description_text.insert("1.0", job_info['description']))
                    self.root.after(0, lambda: self.update_status(f"已抓取：{job_info['title']}"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror(self.texts['error'], str(e)))
                self.root.after(0, lambda: self.update_status(self.texts['status_ready']))
        
        thread = threading.Thread(target=fetch_worker, daemon=True)
        thread.start()
    
    def on_upload_word_click_init(self):
        """上传Word简历（标签1）"""
        self._upload_word_resume(self.resume_text_init)
    
    def _upload_word_resume(self, text_widget):
        """通用的上传Word简历函数"""
        if not DOCX_AVAILABLE:
            messagebox.showerror(self.texts['error'], 
                               "未安装python-docx库，请运行: pip install python-docx")
            return
        
        file_path = filedialog.askopenfilename(
            title="选择Word简历文件",
            filetypes=[("Word文档", "*.docx"), ("Word 97-2003", "*.doc"), ("所有文件", "*.*")]
        )
        
        if file_path:
            try:
                doc = Document(file_path)
                text_content = []
                
                # 读取所有段落
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        text_content.append(text)
                
                # 读取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                resume_text = "\n".join(text_content)
                
                if not resume_text or len(resume_text.strip()) < 50:
                    messagebox.showerror(self.texts['error'], "Word文件内容为空或过短")
                    return
                
                # 更新当前文本控件
                text_widget.delete("1.0", tk.END)
                text_widget.insert("1.0", resume_text)
                
                # 保存到配置和缓存
                self.config['resume_content'] = resume_text
                self.save_config()
                self.save_resume_cache(resume_text)
                
                messagebox.showinfo(self.texts['success'], "简历已上传并保存")
            except Exception as e:
                messagebox.showerror(self.texts['error'], f"读取文件失败: {str(e)}")
    
    def on_preview_resume_click_init(self):
        """预览简历（标签1）"""
        resume_content = self.resume_text_init.get("1.0", tk.END).strip()
        self._preview_resume(resume_content)
    
    def _preview_resume(self, resume_content):
        """通用的预览简历函数"""
        if not resume_content:
            messagebox.showwarning(self.texts['warning'], "简历内容为空")
            return
        
        # 创建预览窗口
        preview_window = tk.Toplevel(self.root)
        preview_window.title("简历预览" if self.language == "zh" else "Resume Preview")
        preview_window.geometry("600x700")
        
        text_widget = scrolledtext.ScrolledText(preview_window, wrap=tk.WORD, font=("Arial", 11))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text_widget.insert("1.0", resume_content)
        text_widget.config(state=tk.DISABLED)
    
    def on_generate_click(self):
        """生成定制简历"""
        # 检查必要信息
        if not self.config.get('api_key'):
            messagebox.showerror(self.texts['error'], "请先配置API Key")
            return
        
        job_description = self.job_description_text.get("1.0", tk.END).strip()
        if not job_description:
            messagebox.showerror(self.texts['error'], "请输入岗位描述")
            return
        
        # 从标签1的简历区域或配置中读取简历
        if hasattr(self, 'resume_text_init'):
            original_resume = self.resume_text_init.get("1.0", tk.END).strip()
        else:
            original_resume = self.config.get('resume_content', '').strip()
        
        if not original_resume:
            messagebox.showerror(self.texts['error'], "请先在「初始化配置」标签页上传或输入原始简历")
            return
        
        resume_language = self.resume_language_var.get()
        
        self.update_status("正在生成定制简历...")
        self.result_text.delete("1.0", tk.END)
        self.result_text.insert("1.0", "正在生成，请稍候...")
        
        def generate_worker():
            try:
                custom_resume, error = self.generate_custom_resume(
                    job_description, original_resume, resume_language
                )
                
                if error:
                    self.root.after(0, lambda: messagebox.showerror(self.texts['error'], error))
                    self.root.after(0, lambda: self.result_text.delete("1.0", tk.END))
                    self.root.after(0, lambda: self.update_status(self.texts['status_ready']))
                else:
                    # 计算匹配度
                    match_score = self.calculate_match_score(job_description, original_resume)
                    
                    # 更新UI
                    self.root.after(0, lambda: self.result_text.delete("1.0", tk.END))
                    self.root.after(0, lambda: self.result_text.insert("1.0", custom_resume))
                    self.root.after(0, lambda: self.update_status(f"生成完成，匹配度: {match_score}%"))
                    self.root.after(0, lambda: messagebox.showinfo(
                        self.texts['success'], 
                        f"简历生成成功！\n匹配度: {match_score}%"
                    ))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror(self.texts['error'], str(e)))
                self.root.after(0, lambda: self.update_status(self.texts['status_ready']))
        
        thread = threading.Thread(target=generate_worker, daemon=True)
        thread.start()
    
    def on_export_pdf_click(self):
        """导出PDF"""
        resume_content = self.result_text.get("1.0", tk.END).strip()
        if not resume_content:
            messagebox.showwarning(self.texts['warning'], "没有可导出的简历内容")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="保存PDF文件",
            defaultextension=".pdf",
            filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        )
        
        if file_path:
            try:
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                
                # 设置中文字体（需要支持中文的字体）
                try:
                    pdf.add_font('SimSun', '', 'simsun.ttc', uni=True)
                    pdf.set_font('SimSun', '', 12)
                except:
                    # 如果没有中文字体，使用默认字体
                    pdf.set_font('Arial', '', 12)
                
                # 添加内容（处理换行）
                lines = resume_content.split('\n')
                for line in lines:
                    if line.strip():
                        pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
                    else:
                        pdf.ln(5)
                
                pdf.output(file_path)
                messagebox.showinfo(self.texts['success'], "PDF已导出")
            except Exception as e:
                messagebox.showerror(self.texts['error'], f"导出失败: {str(e)}")
    
    def on_clear_click(self):
        """清空"""
        self.job_url_entry.delete(0, tk.END)
        self.job_description_text.delete("1.0", tk.END)
        self.result_text.delete("1.0", tk.END)
    
    def create_tab_records(self):
        """创建标签4：投递记录查询"""
        # 顶部提示
        hint_frame = ttk.Frame(self.tab_records)
        hint_frame.pack(fill=tk.X, pady=(0, 15))
        self.hint_records_label = ttk.Label(hint_frame, text=self.texts['hint_records'], 
                              foreground="blue", font=("Arial", 10, "bold"), wraplength=900)
        self.hint_records_label.pack()
        
        # 按钮区域
        button_frame = ttk.Frame(self.tab_records)
        button_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.refresh_btn = ttk.Button(button_frame, text=self.texts['button_refresh'],
                                 command=self.refresh_records)
        self.refresh_btn.pack(side=tk.LEFT, padx=5)
        
        self.export_btn = tk.Button(button_frame, text=self.texts['button_export'],
                               command=self.export_records,
                               bg="#4A90E2", fg="white", font=("Arial", 10, "bold"),
                               padx=15, pady=5, cursor="hand2")
        self.export_btn.pack(side=tk.LEFT, padx=5)
        
        # 记录表格
        self.tree_frame = ttk.LabelFrame(self.tab_records, text=self.texts['frame_records_title'], padding="10")
        self.tree_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建Treeview显示记录
        columns = ('岗位名称', '公司', '投递日期', '匹配度', '状态')
        self.records_tree = ttk.Treeview(self.tree_frame, columns=columns, show='headings', height=20)
        
        # 设置列标题和宽度
        column_widths = {'岗位名称': 250, '公司': 200, '投递日期': 150, '匹配度': 100, '状态': 120}
        for col in columns:
            self.records_tree.heading(col, text=col)
            self.records_tree.column(col, width=column_widths.get(col, 150))
        
        # 滚动条
        scrollbar = ttk.Scrollbar(self.tree_frame, orient=tk.VERTICAL, 
                                  command=self.records_tree.yview)
        self.records_tree.configure(yscrollcommand=scrollbar.set)
        
        # 布局
        self.records_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 初始加载记录
        self.refresh_records()
    
    def refresh_records(self):
        """刷新投递记录"""
        # 清空现有记录
        for item in self.records_tree.get_children():
            self.records_tree.delete(item)
        
        # 从文件加载记录
        records_file = "application_records.json"
        if os.path.exists(records_file):
            try:
                with open(records_file, 'r', encoding='utf-8') as f:
                    records = json.load(f)
                    for record in records:
                        self.records_tree.insert('', tk.END, values=(
                            record.get('job_title', ''),
                            record.get('company', ''),
                            record.get('apply_date', ''),
                            record.get('match_score', ''),
                            record.get('status', '')
                        ))
            except Exception as e:
                print(f"加载记录失败: {e}")
    
    def export_records(self):
        """导出投递记录"""
        records_file = "application_records.json"
        if not os.path.exists(records_file):
            messagebox.showwarning(self.texts['warning'], "没有投递记录")
            return
        
        try:
            with open(records_file, 'r', encoding='utf-8') as f:
                records = json.load(f)
            
            if not records:
                messagebox.showwarning(self.texts['warning'], "没有投递记录")
                return
            
            # 选择导出格式
            file_path = filedialog.asksaveasfilename(
                title="导出投递记录",
                defaultextension=".csv",
                filetypes=[("CSV文件", "*.csv"), ("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
            )
            
            if file_path:
                df = pd.DataFrame(records)
                
                if file_path.endswith('.xlsx'):
                    df.to_excel(file_path, index=False, engine='openpyxl')
                else:
                    df.to_csv(file_path, index=False, encoding='utf-8-sig')
                
                messagebox.showinfo(self.texts['success'], "记录已导出")
        except Exception as e:
            messagebox.showerror(self.texts['error'], f"导出失败: {str(e)}")
    
    def setup_auto_save(self):
        """设置自动保存"""
        # 绑定输入框变化事件，自动保存配置
        self.api_key_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
        self.chrome_dir_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
        self.chrome_profile_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
        self.user_name_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
        self.user_email_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
        self.user_phone_entry.bind('<FocusOut>', lambda e: self.auto_save_config())
    
    def auto_save_config(self):
        """自动保存配置"""
        # 不保存API Key，使用内置的API Key
        # self.config['api_key'] = self.api_key_entry.get()  # 注释掉
        self.config['chrome_user_data_dir'] = self.chrome_dir_entry.get()
        self.config['chrome_profile'] = self.chrome_profile_entry.get()
        self.config['user_name'] = self.user_name_entry.get()
        self.config['user_email'] = self.user_email_entry.get()
        self.config['user_phone'] = self.user_phone_entry.get()
        if hasattr(self, 'expected_salary_var'):
            self.config['expected_salary'] = self.expected_salary_var.get()
        # 保存搜索条件
        if hasattr(self, 'region_var'):
            self.config['region'] = self.region_var.get()
        if hasattr(self, 'search_keyword_entry'):
            self.config['search_keyword'] = self.search_keyword_entry.get()
        if hasattr(self, 'category_var'):
            self.config['job_category'] = self.category_var.get()
        if hasattr(self, 'search_location_entry'):
            self.config['search_location'] = self.search_location_entry.get()
        if hasattr(self, 'match_threshold_entry'):
            try:
                self.config['match_threshold'] = int(self.match_threshold_entry.get())
            except:
                pass
        # 保存投递控制设置
        if hasattr(self, 'max_apply_count_entry'):
            try:
                max_count = int(self.max_apply_count_entry.get())
                if max_count > 15:
                    max_count = 15
                self.config['max_apply_count'] = max_count
            except:
                self.config['max_apply_count'] = 15
        if hasattr(self, 'apply_interval_min_entry'):
            try:
                min_interval = int(self.apply_interval_min_entry.get())
                if min_interval < 5:
                    min_interval = 5
                self.config['apply_interval_min'] = min_interval
            except:
                self.config['apply_interval_min'] = 6
        if hasattr(self, 'apply_interval_max_entry'):
            try:
                max_interval = int(self.apply_interval_max_entry.get())
                if max_interval < 10:
                    max_interval = 10
                self.config['apply_interval_max'] = max_interval
            except:
                self.config['apply_interval_max'] = 12
        # 保存简历内容（从标签1的简历区域）
        if hasattr(self, 'resume_text_init'):
            resume_content = self.resume_text_init.get("1.0", tk.END).strip()
            if resume_content:
                self.config['resume_content'] = resume_content
                self.save_resume_cache(resume_content)
        self.config['language'] = self.language
        self.save_config()
    
    def toggle_language(self):
        """切换中英文 - 全局更新所有界面文字"""
        self.language = "en" if self.language == "zh" else "zh"
        self.texts = self.get_texts(self.language)
        self.config['language'] = self.language
        self.save_config()
        
        # 更新窗口标题
        self.root.title(self.texts['app_title'])
        
        # 更新状态栏
        self.status_label.config(text=self.texts['status_ready'])
        
        # 更新语言切换按钮
        self.lang_button.config(text="EN" if self.language == "zh" else "中文")
        
        # 更新标签页标题
        self.notebook.tab(0, text=self.texts['tab_init'])
        self.notebook.tab(1, text=self.texts['tab_auto'])
        self.notebook.tab(2, text=self.texts['tab_manual'])
        self.notebook.tab(3, text=self.texts['tab_records'])
        
        # 更新标签1（初始化配置）的所有文字
        if hasattr(self, 'hint_init_label'):
            self.hint_init_label.config(text=self.texts['hint_init'])
        if hasattr(self, 'api_frame'):
            self.api_frame.config(text=self.texts['frame_api_title'])
        if hasattr(self, 'browser_frame'):
            self.browser_frame.config(text=self.texts['frame_browser_title'])
        if hasattr(self, 'browser_hint_label'):
            self.browser_hint_label.config(text=self.texts['hint_browser_effect'])
        if hasattr(self, 'browser_detail_label'):
            self.browser_detail_label.config(text=self.texts['hint_browser_detail'])
        if hasattr(self, 'chrome_path_label'):
            self.chrome_path_label.config(text=self.texts['label_chrome_path'])
        if hasattr(self, 'browse_button'):
            self.browse_button.config(text=self.texts['button_browse'])
        if hasattr(self, 'chrome_profile_label'):
            self.chrome_profile_label.config(text=self.texts['label_chrome_profile_name'])
        if hasattr(self, 'chrome_profile_hint_label'):
            self.chrome_profile_hint_label.config(text=self.texts['label_chrome_profile_hint'])
        if hasattr(self, 'user_frame'):
            self.user_frame.config(text=self.texts['frame_user_title'])
        if hasattr(self, 'save_button'):
            self.save_button.config(text=self.texts['button_save'])
        if hasattr(self, 'clear_cache_btn'):
            self.clear_cache_btn.config(text=self.texts['button_clear_cache'])
        
        # 更新标签2（全自动求职）的文字
        if hasattr(self, 'hint_auto_label'):
            self.hint_auto_label.config(text=self.texts['hint_auto_steps'])
        if hasattr(self, 'search_frame'):
            self.search_frame.config(text="🔍 搜索条件" if self.language == "zh" else "🔍 Search Criteria")
        if hasattr(self, 'start_auto_btn'):
            self.start_auto_btn.config(text=self.texts['button_start_auto'])
        if hasattr(self, 'pause_button'):
            self.pause_button.config(text=self.texts['button_pause'])
        if hasattr(self, 'result_frame'):
            self.result_frame.config(text=self.texts['frame_auto_result_title'])
        
        # 更新标签3（手动单岗处理）的文字
        if hasattr(self, 'hint_manual_label'):
            self.hint_manual_label.config(text=self.texts['hint_auto_usage'])
        if hasattr(self, 'job_frame'):
            self.job_frame.config(text=self.texts['frame_job_title'])
        if hasattr(self, 'resume_frame'):
            self.resume_frame.config(text=self.texts['frame_resume_title'])
        if hasattr(self, 'result_frame_manual'):
            self.result_frame_manual.config(text=self.texts['frame_result_title'])
        
        # 更新标签4（投递记录查询）的文字
        if hasattr(self, 'hint_records_label'):
            self.hint_records_label.config(text=self.texts['hint_records'])
        if hasattr(self, 'refresh_btn'):
            self.refresh_btn.config(text=self.texts['button_refresh'])
        if hasattr(self, 'export_btn'):
            self.export_btn.config(text=self.texts['button_export'])
        if hasattr(self, 'tree_frame'):
            self.tree_frame.config(text=self.texts['frame_records_title'])
        
        messagebox.showinfo(self.texts['success'], 
                          "语言已切换" if self.language == "zh" else "Language switched")
    
    # ========== 核心功能函数 ==========
    
    def generate_custom_resume(self, job_description, original_resume, resume_language="auto"):
        """使用DeepSeek API生成定制简历（支持代理服务器）"""
        # 检查是否使用代理服务器
        use_proxy = self.config.get('use_proxy', False)
        
        if use_proxy:
            # 使用代理服务器
            return self._generate_via_proxy(job_description, original_resume, resume_language)
        else:
            # 直接调用API
            return self._generate_direct_api(job_description, original_resume, resume_language)
    
    def _generate_via_proxy(self, job_description, original_resume, resume_language="auto"):
        """通过代理服务器生成简历"""
        proxy_url = self.config.get('proxy_url', 'http://localhost:5000')
        server_api_key = self.config.get('server_api_key', '')
        
        # 检测简历语言
        if resume_language == "auto":
            chinese_chars = len([c for c in original_resume if '\u4e00' <= c <= '\u9fff'])
            resume_language = "zh" if chinese_chars / max(len(original_resume), 1) > 0.3 else "en"
        
        # 构建Prompt
        if resume_language == "zh":
            prompt = f"""你是一位专业的求职顾问。请根据下面的【岗位描述】，重写我的【原始简历】，突出与岗位最匹配的技能和经验。

要求：
1. 保持简历的专业性和真实性
2. 突出与岗位要求最相关的经验和技能
3. 使用专业、简洁的语言
4. 保持简历结构清晰，不要超过一页
5. 保留原始简历中的关键信息（姓名、联系方式、教育背景等）

【岗位描述】
{job_description}

【原始简历】
{original_resume}

请生成定制后的简历："""
        else:
            prompt = f"""You are a professional career consultant. Please rewrite my original resume based on the job description below, highlighting the skills and experiences that best match the position.

Requirements:
1. Maintain professionalism and authenticity
2. Highlight the most relevant experiences and skills for the job requirements
3. Use professional and concise language
4. Keep the resume structure clear, not exceeding one page
5. Retain key information from the original resume (name, contact, education, etc.)

【Job Description】
{job_description}

【Original Resume】
{original_resume}

Please generate the customized resume:"""
        
        try:
            # 调用代理服务器
            headers = {
                "Content-Type": "application/json"
            }
            
            # 如果配置了服务器API Key，添加到请求头
            if server_api_key:
                headers["Authorization"] = f"Bearer {server_api_key}"
            
            data = {
                "model": "deepseek-chat",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 2000
            }
            
            proxy_endpoint = f"{proxy_url.rstrip('/')}/api/chat"
            response = requests.post(proxy_endpoint, json=data, headers=headers, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                generated_resume = result['choices'][0]['message']['content']
                return generated_resume, None
            else:
                return None, "代理服务器返回格式错误"
                
        except requests.exceptions.RequestException as e:
            return None, f"代理服务器请求失败: {str(e)}"
        except Exception as e:
            return None, f"生成失败: {str(e)}"
    
    def _generate_direct_api(self, job_description, original_resume, resume_language="auto"):
        """直接调用DeepSeek API生成简历"""
        # 从配置中获取API Key，如果没有则从GUI输入框获取
        api_key = self.config.get('api_key', '')
        if not api_key and hasattr(self, 'api_key_entry'):
            api_key = self.api_key_entry.get().strip()
        if not api_key:
            return None, "API Key未配置"
        
        # 检测简历语言
        if resume_language == "auto":
            # 简单检测：如果中文字符超过30%，认为是中文简历
            chinese_chars = len([c for c in original_resume if '\u4e00' <= c <= '\u9fff'])
            resume_language = "zh" if chinese_chars / max(len(original_resume), 1) > 0.3 else "en"
        
        # 构建Prompt
        if resume_language == "zh":
            prompt = f"""你是一位专业的求职顾问。请根据下面的【岗位描述】，重写我的【原始简历】，突出与岗位最匹配的技能和经验。

要求：
1. 保持简历的专业性和真实性
2. 突出与岗位要求最相关的经验和技能
3. 使用专业、简洁的语言
4. 保持简历结构清晰，不要超过一页
5. 保留原始简历中的关键信息（姓名、联系方式、教育背景等）

【岗位描述】
{job_description}

【原始简历】
{original_resume}

请生成定制后的简历："""
        else:
            prompt = f"""You are a professional career consultant. Please rewrite my original resume based on the job description below, highlighting the skills and experiences that best match the position.

Requirements:
1. Maintain professionalism and authenticity
2. Highlight the most relevant experiences and skills for the job requirements
3. Use professional and concise language
4. Keep the resume structure clear, not exceeding one page
5. Retain key information from the original resume (name, contact, education, etc.)

【Job Description】
{job_description}

【Original Resume】
{original_resume}

Please generate the customized resume:"""
        
        try:
            url = "https://api.deepseek.com/v1/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": "deepseek-chat",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 2000
            }
            
            response = requests.post(url, json=data, headers=headers, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                generated_resume = result['choices'][0]['message']['content']
                return generated_resume, None
            else:
                return None, "API返回格式错误"
                
        except requests.exceptions.RequestException as e:
            return None, f"网络错误: {str(e)}"
        except Exception as e:
            return None, f"生成失败: {str(e)}"
    
    def fetch_job_info(self, job_url):
        """抓取岗位信息"""
        try:
            # 根据地区确定JobsDB域名
            region = self.config.get('region', '香港 (hk)')
            if 'hk' in region:
                base_url = 'https://hk.jobsdb.com'
            elif 'sg' in region:
                base_url = 'https://sg.jobsdb.com'
            elif 'my' in region:
                base_url = 'https://my.jobsdb.com'
            elif 'ph' in region:
                base_url = 'https://ph.jobsdb.com'
            else:
                base_url = 'https://hk.jobsdb.com'
            
            # 如果URL不完整，补全
            if not job_url.startswith('http'):
                job_url = urljoin(base_url, job_url)
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(job_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 尝试提取岗位描述（JobsDB的HTML结构可能变化，这里提供基础版本）
            job_title = ""
            job_description = ""
            
            # 查找标题
            title_elem = soup.find('h1') or soup.find('title')
            if title_elem:
                job_title = title_elem.get_text(strip=True)
            
            # 查找描述（常见的选择器）
            desc_selectors = [
                'div[data-automation="jobDescription"]',
                '.job-description',
                '#jobDescription',
                'div.jobDescription'
            ]
            
            for selector in desc_selectors:
                desc_elem = soup.select_one(selector)
                if desc_elem:
                    job_description = desc_elem.get_text(strip=True)
                    break
            
            # 如果没找到，尝试查找包含"description"的div
            if not job_description:
                for div in soup.find_all('div', class_=lambda x: x and 'description' in x.lower()):
                    job_description = div.get_text(strip=True)
                    if len(job_description) > 100:
                        break
            
            return {
                'title': job_title,
                'description': job_description,
                'url': job_url
            }, None
            
        except Exception as e:
            return None, f"抓取失败: {str(e)}"
    
    def get_chrome_driver(self, check_running=True):
        """
        获取Chrome浏览器驱动（支持多账号切换和反检测）
        
        参数:
            check_running: 是否检查Chrome是否正在运行（默认True）
        """
        # 前置检查：如果Chrome正在运行，提示用户关闭
        if check_running:
            is_running, process_count = check_chrome_running()
            if is_running:
                msg = f"检测到Chrome浏览器正在运行（{process_count}个进程）。\n\n为了确保使用正确的配置文件，请先关闭所有Chrome窗口。\n\n是否继续？（可能会使用错误的配置文件）" if self.language == "zh" else f"Chrome is running ({process_count} processes).\n\nPlease close all Chrome windows to ensure the correct profile is used.\n\nContinue anyway? (May use wrong profile)"
                if messagebox.askyesno("Chrome正在运行" if self.language == "zh" else "Chrome Running", msg):
                    # 用户选择继续，但给出警告
                    pass
                else:
                    return None, "用户取消了操作，请先关闭Chrome浏览器"
        
        # 优先尝试复用已打开的Chrome浏览器（通过远程调试端口）
        if self.chrome_driver is not None:
            try:
                # 检查driver是否仍然有效
                self.chrome_driver.current_url
                return self.chrome_driver, None
            except:
                # driver已失效，重置
                self.chrome_driver = None
        
        # 尝试连接到已存在的Chrome实例（通过远程调试端口）
        try:
            chrome_options = Options()
            chrome_options.add_experimental_option("debuggerAddress", f"127.0.0.1:{self.chrome_debug_port}")
            driver = webdriver.Chrome(options=chrome_options)
            # 连接成功，保存并返回
            self.chrome_driver = driver
            return driver, None
        except:
            # 连接失败，启动新的Chrome实例
            pass
        
        # 获取用户数据目录和配置文件配置
        user_data_dir = self.config.get('chrome_user_data_dir', '').strip()
        profile_name = self.config.get('chrome_profile', 'Default').strip()
        
        # 如果配置了用户数据目录，必须使用它（确保使用正确的账号）
        if not user_data_dir or not os.path.exists(user_data_dir):
            error_msg = "未配置Chrome用户数据目录或路径不存在。\n\n请在「初始化配置」标签页中设置Chrome用户数据目录路径。" if self.language == "zh" else "Chrome user data directory not configured or path does not exist.\n\nPlease set Chrome user data directory in the initialization configuration tab."
            return None, error_msg
        
        # 验证配置文件是否存在
        profile_path = os.path.join(user_data_dir, profile_name)
        if not os.path.exists(profile_path):
            error_msg = f"配置文件 '{profile_name}' 不存在。\n\n请检查配置文件名称是否正确，或点击「刷新」按钮更新配置文件列表。" if self.language == "zh" else f"Profile '{profile_name}' does not exist.\n\nPlease check the profile name or click 'Refresh' to update the profile list."
            return None, error_msg
        
        # 构建Chrome选项（增强反检测）
        chrome_options = Options()
        
        # 基础稳定性选项
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--disable-software-rasterizer')
        
        # 反检测选项（关键：去除自动化特征）
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation", "enable-logging"])
        chrome_options.add_experimental_option('useAutomationExtension', False)
        
        # 设置用户数据目录和配置文件
        chrome_options.add_argument(f'--user-data-dir={user_data_dir}')
        chrome_options.add_argument(f'--profile-directory={profile_name}')
        
        # 远程调试端口（用于复用浏览器）
        chrome_options.add_argument(f'--remote-debugging-port={self.chrome_debug_port}')
        
        # 设置真实的用户代理
        chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36')
        
        # 其他反检测选项
        chrome_options.add_argument('--disable-infobars')
        chrome_options.add_argument('--disable-notifications')
        chrome_options.add_argument('--disable-popup-blocking')
        
        try:
            # 启动Chrome
            if WEBDRIVER_MANAGER_AVAILABLE:
                service = Service(ChromeDriverManager().install())
                driver = webdriver.Chrome(service=service, options=chrome_options)
            else:
                driver = webdriver.Chrome(options=chrome_options)
            
            # 执行JavaScript去除webdriver特征
            driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {
                'source': '''
                    Object.defineProperty(navigator, 'webdriver', {
                        get: () => undefined
                    });
                    window.navigator.chrome = {
                        runtime: {}
                    };
                    Object.defineProperty(navigator, 'plugins', {
                        get: () => [1, 2, 3, 4, 5]
                    });
                    Object.defineProperty(navigator, 'languages', {
                        get: () => ['en-US', 'en']
                    });
                '''
            })
            
            # 保存driver以便后续复用
            self.chrome_driver = driver
            return driver, None
            
        except Exception as e:
            error_msg = f"启动Chrome失败。\n\n详细错误：{str(e)[:300]}\n\n建议解决方案：\n1. 确保已关闭所有Chrome窗口\n2. 检查Chrome用户数据目录路径是否正确\n3. 检查配置文件名称是否正确\n4. 以管理员身份运行程序（Windows）" if self.language == "zh" else f"Failed to start Chrome.\n\nError: {str(e)[:300]}\n\nSolutions:\n1. Ensure all Chrome windows are closed\n2. Check Chrome user data directory path\n3. Check profile name\n4. Run as administrator (Windows)"
            return None, error_msg
    
    def generate_cover_letter(self, job_description, job_title, company_name, original_resume):
        """使用DeepSeek API生成针对性的cover letter"""
        api_key = self.config.get('api_key', '')
        if not api_key and hasattr(self, 'api_key_entry'):
            api_key = self.api_key_entry.get().strip()
        
        if not api_key:
            return None, "API Key未配置"
        
        # 检测语言
        chinese_chars = len([c for c in job_description if '\u4e00' <= c <= '\u9fff'])
        is_chinese = chinese_chars / max(len(job_description), 1) > 0.3
        
        if is_chinese:
            prompt = f"""你是一位专业的求职顾问。请根据以下信息，为这个岗位写一份专业的求职信（Cover Letter）。

要求：
1. 简洁专业，不超过300字
2. 突出申请人的相关技能和经验
3. 表达对岗位和公司的兴趣
4. 使用正式、礼貌的语言
5. 开头称呼使用"Dear Hiring Manager,"，结尾使用"Sincerely,"

【岗位标题】
{job_title}

【公司名称】
{company_name}

【岗位描述】
{job_description[:1500]}

【申请人简历】
{original_resume[:1000]}

请生成求职信："""
        else:
            prompt = f"""You are a professional career consultant. Please write a professional cover letter for this job position based on the following information.

Requirements:
1. Concise and professional, not exceeding 300 words
2. Highlight the applicant's relevant skills and experience
3. Express interest in the position and company
4. Use formal and polite language
5. Start with "Dear Hiring Manager," and end with "Sincerely,"

【Job Title】
{job_title}

【Company Name】
{company_name}

【Job Description】
{job_description[:1500]}

【Applicant Resume】
{original_resume[:1000]}

Please generate the cover letter:"""
        
        try:
            url = "https://api.deepseek.com/v1/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 800
            }
            
            response = requests.post(url, json=data, headers=headers, timeout=60)
            response.raise_for_status()
            result = response.json()
            
            if 'choices' in result and len(result['choices']) > 0:
                cover_letter = result['choices'][0]['message']['content']
                return cover_letter, None
            else:
                return None, "API返回格式错误"
                
        except Exception as e:
            return None, f"生成失败: {str(e)}"
    
    def convert_resume_to_pdf(self, resume_content, output_path=None):
        """将简历内容转换为PDF文件"""
        if output_path is None:
            output_path = "resume_temp.pdf"
        
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # 设置中文字体
            try:
                # 尝试使用系统字体
                pdf.add_font('SimSun', '', 'simsun.ttc', uni=True)
                pdf.set_font('SimSun', '', 12)
            except:
                try:
                    pdf.add_font('Arial', '', 'arial.ttf', uni=True)
                    pdf.set_font('Arial', '', 12)
                except:
                    pdf.set_font('Arial', '', 12)
            
            # 添加内容
            lines = resume_content.split('\n')
            for line in lines:
                if line.strip():
                    # 处理中文字符
                    try:
                        pdf.cell(0, 10, line, ln=1)
                    except:
                        # 如果编码失败，使用替代方法
                        pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
                else:
                    pdf.ln(5)
            
            pdf.output(output_path)
            return output_path, None
        except Exception as e:
            return None, f"PDF转换失败: {str(e)}"
    
    def auto_apply_job(self, job_url, custom_resume, cover_letter, user_info, resume_pdf_path=None):
        """自动投递岗位 - 完整的JobsDB申请流程"""
        driver, error_or_warning = self.get_chrome_driver()
        if driver is None:
            # 这是真正的错误
            return False, error_or_warning
        elif error_or_warning:
            # 这是警告（成功启动但未使用用户数据目录），记录但继续执行
            print(f"警告: {error_or_warning}")
        
        try:
            # 步骤1: 打开岗位详情页，然后跳转到申请页
            driver.get(job_url)
            time.sleep(5)  # 等待页面加载
            
            # 如果当前是岗位详情页，尝试点击"Apply"或"申请"按钮
            try:
                apply_buttons = [
                    "//button[contains(text(), 'Apply') or contains(text(), '申请')]",
                    "//a[contains(text(), 'Apply') or contains(text(), '申请')]",
                    "//button[contains(@class, 'apply')]",
                    "//a[contains(@class, 'apply')]"
                ]
                for button_xpath in apply_buttons:
                    try:
                        apply_btn = driver.find_element(By.XPATH, button_xpath)
                        driver.execute_script("arguments[0].scrollIntoView(true);", apply_btn)
                        time.sleep(2)
                        apply_btn.click()
                        time.sleep(5)  # 等待跳转到申请页面
                        break
                    except:
                        continue
            except:
                # 如果找不到申请按钮，尝试直接构建申请URL
                if '/job/' in job_url and '/apply/' not in job_url:
                    apply_url = job_url.replace('/job/', '/apply/')
                    driver.get(apply_url)
                    time.sleep(5)
            
            # 步骤2: 处理简历上传（Choose documents步骤）
            # 查找"Upload a resumé"选项或"Select a resumé"选项
            try:
                # 尝试找到上传简历的选项
                upload_resume_radio = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, "//label[contains(text(), 'Upload a resumé') or contains(text(), '上传简历')]"))
                )
                upload_resume_radio.click()
                time.sleep(2)
                
                # 查找文件上传输入框
                file_input = driver.find_element(By.CSS_SELECTOR, "input[type='file']")
                if resume_pdf_path and os.path.exists(resume_pdf_path):
                    file_input.send_keys(os.path.abspath(resume_pdf_path))
                    time.sleep(3)
            except:
                # 如果上传失败，尝试选择已有简历
                try:
                    select_resume_radio = driver.find_element(By.XPATH, "//label[contains(text(), 'Select a resumé') or contains(text(), '选择简历')]")
                    select_resume_radio.click()
                    time.sleep(2)
                except:
                    pass
            
            # 点击Continue按钮进入下一步
            try:
                continue_btn = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Continue') or contains(text(), '继续')]"))
                )
                continue_btn.click()
                time.sleep(5)
            except:
                pass
            
            # 步骤3: 填写Cover Letter（Answer employer questions步骤）
            try:
                # 查找"Write a cover letter"选项
                write_cover_radio = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, "//label[contains(text(), 'Write a cover letter') or contains(text(), '写求职信')]"))
                )
                write_cover_radio.click()
                time.sleep(2)
                
                # 查找cover letter文本区域
                cover_letter_selectors = [
                    "textarea[name*='cover']",
                    "textarea[id*='cover']",
                    "textarea[placeholder*='cover' i]",
                    "textarea[placeholder*='letter' i]"
                ]
                for selector in cover_letter_selectors:
                    try:
                        cover_textarea = driver.find_element(By.CSS_SELECTOR, selector)
                        driver.execute_script("arguments[0].scrollIntoView(true);", cover_textarea)
                        time.sleep(1)
                        cover_textarea.clear()
                        cover_textarea.send_keys(cover_letter)
                        break
                    except:
                        continue
            except Exception as e:
                return False, f"填写Cover Letter失败: {str(e)}"
            
            # 填写期望薪资（如果有这个字段）
            expected_salary = user_info.get('expected_salary', '$20K')
            try:
                # 查找期望薪资输入框或下拉框
                salary_selectors = [
                    "input[name*='salary']",
                    "input[id*='salary']",
                    "select[name*='salary']",
                    "select[id*='salary']"
                ]
                for selector in salary_selectors:
                    try:
                        salary_element = driver.find_element(By.CSS_SELECTOR, selector)
                        if salary_element.tag_name == 'select':
                            from selenium.webdriver.support.ui import Select
                            select = Select(salary_element)
                            # 尝试选择匹配的选项
                            for option in select.options:
                                if expected_salary in option.text:
                                    select.select_by_visible_text(option.text)
                                    break
                        else:
                            salary_element.clear()
                            salary_element.send_keys(expected_salary)
                        break
                    except:
                        continue
            except:
                pass  # 如果找不到薪资字段，继续执行
            
            # 回答其他雇主问题（如工作权限等）
            try:
                # 查找工作权限相关的单选按钮
                work_rights_selectors = [
                    "//label[contains(text(), 'Hong Kong SAR citizen')]",
                    "//label[contains(text(), '香港永久居民')]",
                    "//input[@value='citizen' or @value='permanent']"
                ]
                for selector in work_rights_selectors:
                    try:
                        if selector.startswith("//"):
                            element = driver.find_element(By.XPATH, selector)
                        else:
                            element = driver.find_element(By.CSS_SELECTOR, selector)
                        element.click()
                        time.sleep(1)
                        break
                    except:
                        continue
            except:
                pass  # 如果找不到，继续执行
            
            # 点击Continue进入下一步
            try:
                continue_btn = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Continue') or contains(text(), '继续')]"))
                )
                continue_btn.click()
                time.sleep(5)
            except:
                pass
            
            # 步骤4: 最终提交（Review and submit步骤）
            try:
                # 查找Submit application按钮
                submit_btn = WebDriverWait(driver, 15).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Submit application') or contains(text(), '提交申请')]"))
                )
                driver.execute_script("arguments[0].scrollIntoView(true);", submit_btn)
                time.sleep(2)
                submit_btn.click()
                time.sleep(3)
                
                return True, "申请已成功提交"
            except Exception as e:
                return False, f"提交失败: {str(e)}。请手动检查并提交。"
            
        except Exception as e:
            return False, f"自动投递失败: {str(e)}"
        finally:
            # 保持浏览器打开，让用户查看结果
            pass
    
    def calculate_match_score(self, job_description, resume):
        """使用DeepSeek API计算简历与岗位的匹配度（支持代理服务器）"""
        # 检查是否使用代理服务器
        use_proxy = self.config.get('use_proxy', False)
        
        if use_proxy:
            return self._calculate_match_via_proxy(job_description, resume)
        else:
            return self._calculate_match_direct_api(job_description, resume)
    
    def _calculate_match_via_proxy(self, job_description, resume):
        """通过代理服务器计算匹配度"""
        proxy_url = self.config.get('proxy_url', 'http://localhost:5000')
        server_api_key = self.config.get('server_api_key', '')
        
        prompt = f"""你是一位专业的HR顾问。请评估以下简历与岗位描述的匹配度。

要求：
1. 仔细分析岗位描述中的关键要求（技能、经验、学历等）
2. 评估简历中是否包含这些关键要求
3. 给出0-100分的匹配度评分
4. 只输出一个数字（0-100之间的整数），不要输出其他文字

【岗位描述】
{job_description[:2000]}

【简历内容】
{resume[:2000]}

请直接输出匹配度分数（0-100的整数）："""
        
        try:
            headers = {"Content-Type": "application/json"}
            if server_api_key:
                headers["Authorization"] = f"Bearer {server_api_key}"
            
            data = {
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 50
            }
            
            proxy_endpoint = f"{proxy_url.rstrip('/')}/api/chat"
            response = requests.post(proxy_endpoint, json=data, headers=headers, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            if "choices" in result and len(result["choices"]) > 0:
                score_text = result["choices"][0]["message"]["content"].strip()
                score_match = re.search(r'\d+', score_text)
                if score_match:
                    score = int(score_match.group())
                    return max(0, min(100, score))
            
            return self._calculate_match_simple(job_description, resume)
            
        except Exception as e:
            return self._calculate_match_simple(job_description, resume)
    
    def _calculate_match_direct_api(self, job_description, resume):
        """直接调用API计算匹配度"""
        api_key = self.config.get('api_key', '')
        if not api_key and hasattr(self, 'api_key_entry'):
            api_key = self.api_key_entry.get().strip()
        
        if not api_key:
            return self._calculate_match_simple(job_description, resume)
        
        prompt = f"""你是一位专业的HR顾问。请评估以下简历与岗位描述的匹配度。

要求：
1. 仔细分析岗位描述中的关键要求（技能、经验、学历等）
2. 评估简历中是否包含这些关键要求
3. 给出0-100分的匹配度评分
4. 只输出一个数字（0-100之间的整数），不要输出其他文字

【岗位描述】
{job_description[:2000]}

【简历内容】
{resume[:2000]}

请直接输出匹配度分数（0-100的整数）："""
        
        try:
            url = "https://api.deepseek.com/v1/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 50
            }
            
            response = requests.post(url, json=data, headers=headers, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            if "choices" in result and len(result["choices"]) > 0:
                score_text = result["choices"][0]["message"]["content"].strip()
                score_match = re.search(r'\d+', score_text)
                if score_match:
                    score = int(score_match.group())
                    return max(0, min(100, score))
            
            return self._calculate_match_simple(job_description, resume)
            
        except Exception as e:
            return self._calculate_match_simple(job_description, resume)
    
    def _calculate_match_simple(self, job_description, resume):
        """简单关键词匹配（备用方案）"""
        job_keywords = set(re.findall(r'\b\w{4,}\b', job_description.lower()))
        resume_keywords = set(re.findall(r'\b\w{4,}\b', resume.lower()))
        if not job_keywords:
            return 0
        matched = len(job_keywords & resume_keywords)
        match_score = int((matched / len(job_keywords)) * 100)
        return min(match_score, 100)
    
    def get_daily_apply_count(self):
        """获取今日已投递数量"""
        records_file = "application_records.json"
        today = datetime.now().strftime('%Y-%m-%d')
        count = 0
        
        if os.path.exists(records_file):
            try:
                with open(records_file, 'r', encoding='utf-8') as f:
                    records = json.load(f)
                    for record in records:
                        apply_date = record.get('apply_date', '')
                        if apply_date.startswith(today):
                            count += 1
            except:
                pass
        
        return count
    
    def save_application_record(self, job_title, company, job_url, match_score, status="已投递"):
        """保存投递记录"""
        records_file = "application_records.json"
        records = []
        
        if os.path.exists(records_file):
            try:
                with open(records_file, 'r', encoding='utf-8') as f:
                    records = json.load(f)
            except:
                records = []
        
        record = {
            'job_title': job_title,
            'company': company,
            'job_url': job_url,
            'apply_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'match_score': f"{match_score}%",
            'status': status
        }
        
        records.append(record)
        
        try:
            with open(records_file, 'w', encoding='utf-8') as f:
                json.dump(records, f, ensure_ascii=False, indent=2)
            
            # 同时保存到Excel（如果pandas可用）
            try:
                df = pd.DataFrame(records)
                df.to_excel("application_records.xlsx", index=False, engine='openpyxl')
            except:
                pass
            
            return True
        except Exception as e:
            print(f"保存记录失败: {e}")
            return False
    
    # ========== 简历缓存功能 ==========
    
    def load_resume_cache(self):
        """从缓存文件加载简历内容"""
        cache_file = "resume_cache.txt"
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                print(f"加载简历缓存失败: {e}")
        return ""
    
    def save_resume_cache(self, resume_content):
        """保存简历内容到缓存文件"""
        cache_file = "resume_cache.txt"
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                f.write(resume_content)
            return True
        except Exception as e:
            print(f"保存简历缓存失败: {e}")
            return False
    
    # ========== 自动搜索功能 ==========
    
    def scrape_job_urls(self, search_criteria, max_pages=5):
        """
        自动搜索JobsDB并抓取岗位URL列表
        
        参数:
            search_criteria: 搜索条件字典，例如 {'keyword': 'Administrative Officer', 'location': 'Hong Kong'}
            max_pages: 最大抓取页数（默认5页）
        
        返回:
            (成功标志, URL列表或错误信息)
        """
        keyword = search_criteria.get('keyword', '').strip()
        location = search_criteria.get('location', '').strip()
        region = self.config.get('region', '香港 (hk)')
        
        if not keyword:
            return False, "搜索关键词不能为空"
        
        # 确定JobsDB域名
        if 'hk' in region:
            base_url = 'https://hk.jobsdb.com/hk/'
        elif 'sg' in region:
            base_url = 'https://sg.jobsdb.com/'
        elif 'my' in region:
            base_url = 'https://my.jobsdb.com/'
        elif 'ph' in region:
            base_url = 'https://ph.jobsdb.com/'
        else:
            base_url = 'https://hk.jobsdb.com/hk/'
        
        driver = None
        job_urls = set()  # 使用set自动去重
        
        try:
            # 使用统一的get_chrome_driver函数（不检查进程，因为可能已经通过"打开目标网站"打开了）
            driver, error_or_warning = self.get_chrome_driver(check_running=False)
            if driver is None:
                # 这是真正的错误
                return False, error_or_warning
            elif error_or_warning:
                # 这是警告（成功启动但未使用用户数据目录）
                # 记录警告但继续执行
                if hasattr(self, 'log_auto_result'):
                    self.log_auto_result(f"{error_or_warning}\n\n")
            
            driver.maximize_window()
            
            # 访问JobsDB首页
            driver.get(base_url)
            time.sleep(2)
            
            # 输入关键词
            keyword_selectors = [
                (By.ID, "searchKeywordsField"),
                (By.NAME, "searchKeywordsField"),
                (By.CSS_SELECTOR, "input[placeholder*='Job title']"),
                (By.CSS_SELECTOR, "input[placeholder*='关键词']"),
            ]
            
            keyword_input = None
            for by, selector in keyword_selectors:
                try:
                    keyword_input = WebDriverWait(driver, 5).until(
                        EC.presence_of_element_located((by, selector))
                    )
                    break
                except:
                    continue
            
            if not keyword_input:
                return False, "无法找到关键词搜索框"
            
            keyword_input.clear()
            keyword_input.send_keys(keyword)
            time.sleep(1)
            
            # 输入地点
            if location:
                location_selectors = [
                    (By.ID, "searchLocationField"),
                    (By.NAME, "searchLocationField"),
                    (By.CSS_SELECTOR, "input[placeholder*='Location']"),
                ]
                
                for by, selector in location_selectors:
                    try:
                        location_input = driver.find_element(by, selector)
                        location_input.clear()
                        location_input.send_keys(location)
                        time.sleep(1)
                        break
                    except:
                        continue
            
            # 点击搜索
            search_button_selectors = [
                (By.CSS_SELECTOR, "button[type='submit']"),
                (By.XPATH, "//button[contains(text(), 'Search')]"),
                (By.XPATH, "//button[contains(text(), '搜索')]"),
            ]
            
            for by, selector in search_button_selectors:
                try:
                    search_button = driver.find_element(by, selector)
                    search_button.click()
                    break
                except:
                    continue
            
            time.sleep(3)
            
            # 抓取多页结果
            current_page = 1
            while current_page <= max_pages:
                time.sleep(2)
                
                # 查找岗位链接
                job_link_selectors = [
                    (By.CSS_SELECTOR, "a[href*='/hk/en/job/']"),
                    (By.CSS_SELECTOR, "a[href*='/hk/job/']"),
                    (By.XPATH, "//a[contains(@href, '/job/')]"),
                ]
                
                for by, selector in job_link_selectors:
                    try:
                        links = driver.find_elements(by, selector)
                        for link in links:
                            href = link.get_attribute('href')
                            if href and '/job/' in href:
                                if href.startswith('http'):
                                    job_urls.add(href)
                                elif href.startswith('/'):
                                    job_urls.add('https://hk.jobsdb.com' + href)
                        if job_urls:
                            break
                    except:
                        continue
                
                # 尝试点击下一页
                if current_page < max_pages:
                    next_selectors = [
                        (By.XPATH, "//a[contains(text(), 'Next')]"),
                        (By.XPATH, "//a[contains(text(), '下一页')]"),
                        (By.CSS_SELECTOR, "a[aria-label*='Next']"),
                    ]
                    
                    next_button = None
                    for by, selector in next_selectors:
                        try:
                            next_button = driver.find_element(by, selector)
                            if next_button.is_enabled() and next_button.is_displayed():
                                break
                        except:
                            continue
                    
                    if next_button:
                        try:
                            driver.execute_script("arguments[0].click();", next_button)
                            time.sleep(3)
                            current_page += 1
                        except:
                            break
                    else:
                        break
                else:
                    break
            
            return True, list(job_urls)
            
        except Exception as e:
            return False, f"抓取失败: {str(e)}"
        finally:
            if driver:
                driver.quit()


def main():
    """主函数"""
    root = tk.Tk()
    app = ResumeGeneratorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()

